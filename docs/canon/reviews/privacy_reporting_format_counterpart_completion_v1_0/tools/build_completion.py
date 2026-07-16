#!/usr/bin/env python3
"""Build founder-authorized Privacy and Reporting deterministic counterparts."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


REPO = Path(__file__).resolve().parents[5]
REVIEW = REPO / "docs/canon/reviews/privacy_reporting_format_counterpart_completion_v1_0"
SOURCES = REVIEW / "exact_sources"
APPROVED = REPO / "docs/canon/founder_approved_sources"

PRIVACY_SOURCE = Path("/Users/rianray/Downloads/MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_REVIEWED_DRAFT.md")
REPORTING_SOURCE = Path("/Users/rianray/Downloads/MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0.docx")
PRIVACY_SOURCE_SHA = "30858ed46725a8adaf7104193525aa1074dd8421555ded92ede034ec3339202e"
REPORTING_SOURCE_SHA = "d6cc3cc2add501d2d34d947de4e7a2ee7c6af77aa9669606c95514361657504d"

PRIVACY_MD = APPROVED / "MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.md"
PRIVACY_DOCX = APPROVED / "MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.docx"
REPORTING_MD = APPROVED / "MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0_FOUNDER_APPROVED.md"
REPORTING_DOCX = APPROVED / "MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0_FOUNDER_APPROVED.docx"

CLASSIFICATION = "FOUNDER_AUTHORIZED_DETERMINISTIC_FORMAT_COUNTERPART_DERIVED_FROM_EXACT_SUBSTANTIVE_SOURCE"
FINAL_DISPOSITION = "PRIVACY_AND_REPORTING_FOUNDER_APPROVED_FORMAT_COUNTERPARTS_COMPLETE_PENDING_SEPARATE_ADOPTION"

AUTHORITY = {
    "adoption": False,
    "lock": False,
    "governance_v1_baseline_adoption": False,
    "governance_v1_baseline_lock": False,
    "implementation": False,
    "migration": False,
    "runtime": False,
    "production": False,
    "provider_activation": False,
    "deployment": False,
    "customer_data_processing": False,
    "public_launch": False,
    "security_certification": False,
    "public_trust_claims": False,
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def rel(path: Path) -> str:
    return path.relative_to(REPO).as_posix()


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def normalized(value: str) -> str:
    value = value.replace("\u00a0", " ").replace("\u2013", "-").replace("\u2014", "-")
    return re.sub(r"\s+", " ", value).strip()


def plain_markdown(value: str) -> str:
    value = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"(`{1,3}|\*\*|__|\*|_)", "", value)
    return normalized(value)


def iter_blocks(document: Document) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def escape_table_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>").strip()


def docx_to_markdown(path: Path) -> str:
    lines: list[str] = []
    for block in iter_blocks(Document(path)):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                lines.append("")
                continue
            style = (block.style.name or "").lower()
            heading = re.search(r"heading\s*(\d+)", style)
            if heading:
                lines.extend([f"{'#' * min(int(heading.group(1)), 6)} {text.replace(chr(10), ' ')}", ""])
            elif "title" in style:
                lines.extend([f"# {text.replace(chr(10), ' ')}", ""])
            elif "subtitle" in style:
                lines.extend([f"## {text.replace(chr(10), ' ')}", ""])
            elif "list bullet" in style:
                lines.extend([f"- {text}", ""])
            elif "list number" in style:
                lines.extend([f"1. {text}", ""])
            else:
                lines.extend([text, ""])
        else:
            rows = [[escape_table_cell(cell.text) for cell in row.cells] for row in block.rows]
            if not rows:
                continue
            width = max(map(len, rows))
            rows = [row + [""] * (width - len(row)) for row in rows]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
            lines.append("")
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n"


def add_page_number(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip().replace("<br>", "\n").replace("\\|", "|") for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def markdown_to_docx(markdown: str, output: Path, title: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.3)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for level in range(1, 7):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(55, 47, 73)
        style.font.size = Pt(max(11, 22 - level * 2))
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "EQUINESYNC  |  FOUNDER-APPROVED CANON SOURCE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(111, 96, 128)
    add_page_number(section.footer.paragraphs[0])

    lines = markdown.splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = parse_table(table_lines)
            if rows:
                width = max(map(len, rows))
                table = document.add_table(rows=len(rows), cols=width)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for column_index in range(width):
                        cell = table.cell(row_index, column_index)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.text = plain_markdown(row[column_index]) if column_index < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                if row_index == 0:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                        if row_index == 0:
                            set_cell_shading(cell, "40364D")
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            paragraph = document.add_paragraph(plain_markdown(heading.group(2)), style=f"Heading {len(heading.group(1))}")
            if first_heading:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_heading = False
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(plain_markdown(line[2:]))
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.right_indent = Inches(0.3)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(80, 70, 92)
            index += 1
            continue
        if re.match(r"^[-*]\s+", line):
            document.add_paragraph(plain_markdown(re.sub(r"^[-*]\s+", "", line)), style="List Bullet")
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            document.add_paragraph(plain_markdown(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            index += 1
            continue
        document.add_paragraph(plain_markdown(line))
        index += 1

    document.core_properties.title = title
    document.core_properties.subject = CLASSIFICATION
    document.core_properties.author = "EquineSync"
    document.core_properties.keywords = "privacy, data protection, founder approved, deterministic counterpart"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def privacy_successor(source: str) -> str:
    marker = "## Authority Boundary"
    if marker not in source:
        raise ValueError("Privacy authority-boundary marker missing")
    body = marker + source.split(marker, 1)[1]
    control = f"""# MASTER PRIVACY AND DATA PROTECTION MODEL V2.0

**FOUNDER APPROVED - FOUNDER-APPROVED CANON SOURCE - NOT YET ADOPTED - NOT YET LOCKED**

# Document Control

Document: MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.md

Version: 2.0 Founder Approved

Original reviewed-draft date: July 13, 2026

Founder Approval Date: July 16, 2026

Founder Approval Status: FOUNDER APPROVED

Constitutional Status: FOUNDER-APPROVED CANON SOURCE

Adoption Status: NOT YET ADOPTED

Lock Status: NOT YET LOCKED

Implementation Authority: FALSE

Runtime Authority: FALSE

Production Authority: FALSE

Provider Activation Authority: FALSE

Public Launch Authority: FALSE

Public Trust Claim Authority: FALSE

Evidentiary Classification: {CLASSIFICATION}

## Format-Counterpart Provenance

This founder-approved Markdown successor preserves the substantive text of the exact mounted reviewed-draft Markdown source with SHA-256 `{PRIVACY_SOURCE_SHA}`. The accompanying DOCX is a founder-authorized deterministic format counterpart generated from that exact mounted Markdown source. It is not represented as a recovered historical DOCX.

"""
    return control + body.rstrip() + "\n"


def reporting_successor(source_markdown: str) -> str:
    control = f"""# Reporting Format-Counterpart Control

| Field | Value |
|---|---|
| Founder Approval Status | FOUNDER APPROVED |
| Constitutional Status | FOUNDER-APPROVED CANON SOURCE |
| Adoption Status | NOT YET ADOPTED |
| Lock Status | NOT YET LOCKED |
| Implementation Authority | FALSE |
| Runtime Authority | FALSE |
| Production Authority | FALSE |
| Provider Activation Authority | FALSE |
| Public Launch Authority | FALSE |
| Public Trust Claim Authority | FALSE |
| Evidentiary Classification | {CLASSIFICATION} |

## Format-Counterpart Provenance

This Markdown file is a founder-authorized deterministic format counterpart generated from the exact mounted DOCX source with SHA-256 `{REPORTING_SOURCE_SHA}`. It is not represented as a recovered historical Markdown source. The exact mounted DOCX remains the source-of-content authority.

# Exact Substantive Source Content

"""
    return control + source_markdown.rstrip() + "\n"


def docx_units(path: Path) -> dict:
    paragraphs = []
    headings = []
    lists = []
    tables = []
    for block in iter_blocks(Document(path)):
        if isinstance(block, Paragraph):
            text = normalized(block.text)
            if not text:
                continue
            style = block.style.name or ""
            paragraphs.append(text)
            if style.lower().startswith(("heading", "title", "subtitle")):
                headings.append((style, text))
            if style.lower().startswith("list"):
                lists.append(text)
        else:
            tables.append([[normalized(cell.text) for cell in row.cells] for row in block.rows])
    return {"paragraphs": paragraphs, "headings": headings, "lists": lists, "tables": tables}


def markdown_units(text: str) -> dict:
    paragraphs = []
    headings = []
    lists = []
    tables = []
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            group = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                group.append(lines[index].strip())
                index += 1
            tables.append([[plain_markdown(cell) for cell in row] for row in parse_table(group)])
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            value = plain_markdown(heading.group(2))
            paragraphs.append(value)
            headings.append((len(heading.group(1)), value))
        elif re.match(r"^[-*]\s+", line):
            value = plain_markdown(re.sub(r"^[-*]\s+", "", line))
            paragraphs.append(value)
            lists.append(value)
        elif re.match(r"^\d+\.\s+", line):
            value = plain_markdown(re.sub(r"^\d+\.\s+", "", line))
            paragraphs.append(value)
            lists.append(value)
        elif line.startswith("> "):
            paragraphs.append(plain_markdown(line[2:]))
        else:
            paragraphs.append(plain_markdown(line))
        index += 1
    return {"paragraphs": paragraphs, "headings": headings, "lists": lists, "tables": tables}


def identifiers(text: str) -> list[str]:
    return sorted(set(re.findall(r"\b(?:FD-[A-Z0-9]+\d+|[A-Z]{2,12}-\d{2,4}|C0-\d{3})\b", text)))


def defined_terms(text: str) -> list[str]:
    patterns = [r"\*\*([^*]{2,80})\*\*\s+(?:means|includes)", r"\b([A-Z][A-Za-z /-]{2,70})\s+means\b"]
    found = set()
    for pattern in patterns:
        found.update(normalized(value) for value in re.findall(pattern, text))
    return sorted(found)


def authority_signals(text: str) -> dict:
    lowered = text.lower()
    return {
        "shall_not": lowered.count("shall not"),
        "must_not": lowered.count("must not"),
        "may_not": lowered.count("may not"),
        "unauthorized": lowered.count("unauthorized"),
        "false_markers": len(re.findall(r"\bFALSE\b", text)),
    }


def replace_once(path: Path, old: str, new: str) -> None:
    text = path.read_text(encoding="utf-8")
    if new in text:
        return
    if old not in text:
        raise ValueError(f"Expected text not found in {path}: {old[:80]}")
    write(path, text.replace(old, new, 1))


def update_json_rows(path: Path, collection: str, updates: dict[str, dict]) -> None:
    data = json.loads(path.read_text(encoding="utf-8"))
    seen = set()
    for row in data[collection]:
        row_id = row.get("record_id") or row.get("row_id")
        if row_id in updates:
            row.update(updates[row_id])
            seen.add(row_id)
    if seen != set(updates):
        raise ValueError(f"Missing rows in {path}: {set(updates) - seen}")
    write(path, json.dumps(data, indent=2))


def update_companions(privacy_hash: str, reporting_hash: str) -> list[str]:
    report_path = rel(REVIEW / "PRIVACY_REPORTING_FORMAT_COUNTERPART_COMPLETION_REPORT.md")
    source_register = REPO / "docs/canon/reviews/governance_v1_0_c0_lifecycle_completion_batch_1/C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.json"
    state_register = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_CURRENT_LIFECYCLE_STATE_REGISTER.json"
    row_ledger = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.json"
    delta = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_HISTORICAL_TO_CURRENT_DELTA.json"

    common = {
        "current_category": "adoption_and_lock_required",
        "current_exact_source_status": "EXACT_SOURCE_AVAILABLE_FORMAT_COUNTERPART_COMPLETE",
        "founder_status": "FOUNDER_APPROVED",
        "adoption_state": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION",
        "lock_state": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION",
        "lifecycle_evidence": report_path,
        "unresolved_blocker": True,
        "remedy": "Complete separately authorized founder adoption and lock review; no operational authority is implied.",
    }
    update_json_rows(source_register, "rows", {
        "C0-023": {**common, "current_repository_path": rel(PRIVACY_MD), "current_sha256": privacy_hash, "current_controlling_artifact": rel(PRIVACY_MD)},
        "C0-035": {**common, "current_repository_path": rel(REPORTING_DOCX), "current_sha256": reporting_hash, "current_controlling_artifact": rel(REPORTING_DOCX)},
    })
    update_json_rows(state_register, "rows", {
        "C0-023": {**common, "current_repository_path": rel(PRIVACY_MD), "current_sha256": privacy_hash, "current_controlling_artifact": rel(PRIVACY_MD)},
        "C0-035": {**common, "current_repository_path": rel(REPORTING_DOCX), "current_sha256": reporting_hash, "current_controlling_artifact": rel(REPORTING_DOCX)},
    })
    update_json_rows(row_ledger, "records", {
        "C0-023": {
            "current_category": "adoption_and_lock_required", "source": rel(PRIVACY_MD), "expected_c0_sha256": privacy_hash,
            "founder_status": "FOUNDER_APPROVED", "adoption_state": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION",
            "lock_state": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION", "next_step": "Complete separately authorized founder adoption and lock review.",
            "track_d_status": "FORMAT_COUNTERPART_COMPLETE", "track_d_classification": CLASSIFICATION,
        },
        "C0-035": {
            "current_category": "adoption_and_lock_required", "source": rel(REPORTING_DOCX), "expected_c0_sha256": reporting_hash,
            "founder_status": "FOUNDER_APPROVED", "adoption_state": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION",
            "lock_state": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION", "next_step": "Complete separately authorized founder adoption and lock review.",
            "track_d_status": "FORMAT_COUNTERPART_COMPLETE", "track_d_classification": CLASSIFICATION,
        },
    })
    update_json_rows(delta, "rows", {
        "C0-023": {"current_category": "adoption_and_lock_required", "current_founder_status": "FOUNDER_APPROVED", "current_adoption": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION", "current_lock": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION", "current_path": rel(PRIVACY_MD)},
        "C0-035": {"current_category": "adoption_and_lock_required", "current_founder_status": "FOUNDER_APPROVED", "current_adoption": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION", "current_lock": "PENDING_SEPARATE_FOUNDER_AUTHORIZATION", "current_path": rel(REPORTING_DOCX)},
    })

    paths = {
        "batch_md": REPO / "docs/canon/reviews/governance_v1_0_c0_lifecycle_completion_batch_1/C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.md",
        "state_md": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_CURRENT_LIFECYCLE_STATE_REGISTER.md",
        "ledger_md": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.md",
        "blocker_md": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_UNRESOLVED_LIFECYCLE_BLOCKER_LEDGER.md",
        "delta_md": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_HISTORICAL_TO_CURRENT_DELTA.md",
    }
    for row_id, family, artifact in (("C0-023", "Master Privacy and Data Protection Model", rel(PRIVACY_MD)), ("C0-035", "Master Reporting, Analytics, and Business Intelligence Model", rel(REPORTING_DOCX))):
        text = paths["batch_md"].read_text(encoding="utf-8")
        text = re.sub(rf"^\| {row_id} \|.*$", f"| {row_id} | {family} | adoption_and_lock_required | EXACT_SOURCE_AVAILABLE_FORMAT_COUNTERPART_COMPLETE | FOUNDER_APPROVED | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | `{artifact}` |", text, flags=re.M)
        write(paths["batch_md"], text)
        text = paths["state_md"].read_text(encoding="utf-8")
        text = re.sub(rf"^\| {row_id} \|.*$", f"| {row_id} | {family} | EXACT_SOURCE_AVAILABLE_FORMAT_COUNTERPART_COMPLETE | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | `{artifact}` | True |", text, flags=re.M)
        write(paths["state_md"], text)
        text = paths["ledger_md"].read_text(encoding="utf-8")
        text = re.sub(rf"^\| {row_id} \|.*$", f"| {row_id} | TRACK_D_MISSING_SOURCE_RECOVERY | adoption_and_lock_required | FOUNDER_APPROVED | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | Complete separately authorized founder adoption and lock review. |", text, flags=re.M)
        write(paths["ledger_md"], text)
        text = paths["blocker_md"].read_text(encoding="utf-8")
        text = re.sub(rf"^\| {row_id} \|.*$", f"| {row_id} | {family} | adoption_and_lock_required | FOUNDER_APPROVED | Exact source and deterministic format counterpart complete; adoption and lock remain separately gated. |", text, flags=re.M)
        write(paths["blocker_md"], text)
        text = paths["delta_md"].read_text(encoding="utf-8")
        text = re.sub(rf"^\| {row_id} \|.*$", f"| {row_id} | UNRESOLVED | FOUNDER_APPROVED | UNRESOLVED | PENDING_SEPARATE_FOUNDER_AUTHORIZATION | `{artifact}` |", text, flags=re.M)
        write(paths["delta_md"], text)

    completion_note = f"**July 16, 2026 Privacy/Reporting format-counterpart completion:** C0-023 and C0-035 exact substantive sources are available and their founder-authorized deterministic format counterparts are complete. See [`PRIVACY_REPORTING_FORMAT_COUNTERPART_COMPLETION_REPORT.md`](../reviews/privacy_reporting_format_counterpart_completion_v1_0/PRIVACY_REPORTING_FORMAT_COUNTERPART_COMPLETION_REPORT.md). Adoption, lock, implementation, runtime, production, provider, launch, and public-trust authority remain false."
    for path, prior in [
        (REPO / "docs/canon/companions/EQUINESYNC_CONSTITUTIONAL_AUTHORITY_MATRIX_V1_2.md", "**July 16, 2026 lifecycle reconciliation:**"),
        (REPO / "docs/canon/companions/EQUINESYNC_CONSTITUTIONAL_CROSS_REFERENCE_INDEX_V1_2.md", "**July 16, 2026 lifecycle reconciliation:**"),
        (REPO / "docs/canon/companions/EQUINESYNC_CANON_DEPENDENCY_MAP_V1_2.md", "**July 16, 2026 lifecycle reconciliation:**"),
        (REPO / "docs/canon/companions/MASTER_EQUINESYNC_FOUNDER_DECISION_REGISTER_V1_0.md", "**July 16, 2026 founder event:**"),
        (REPO / "docs/canon/companions/MASTER_EQUINESYNC_GOVERNANCE_REQUIREMENT_INDEX_V1_0.md", "**July 16, 2026 lifecycle reconciliation:**"),
        (REPO / "docs/canon/companions/MASTER_EQUINESYNC_REQUIREMENT_TRACEABILITY_MATRIX_V1_0.md", "**July 16, 2026 lifecycle reconciliation:**"),
    ]:
        text = path.read_text(encoding="utf-8")
        if completion_note not in text:
            position = text.find("\n", text.find(prior))
            text = text[:position + 1] + "\n" + completion_note + "\n" + text[position + 1:]
            write(path, text)

    event = {
        "disposition": "FOUNDER_AUTHORIZED_PRIVACY_AND_REPORTING_DETERMINISTIC_FORMAT_COUNTERPART_COMPLETION",
        "source": report_path,
        "rows": ["C0-023", "C0-035"],
        "adoption": False, "lock": False, "implementation": False, "production": False,
    }
    for path in [
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_FOUNDER_DECISION_REGISTER_V1_0.json",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_GOVERNANCE_REQUIREMENT_INDEX_V1_0.json",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_REQUIREMENT_TRACEABILITY_MATRIX_V1_0.json",
    ]:
        data = json.loads(path.read_text(encoding="utf-8"))
        data["privacy_reporting_format_counterpart_completion"] = event
        write(path, json.dumps(data, indent=2))

    canon_index = REPO / "docs/canon/CANON_INDEX.md"
    index_text = canon_index.read_text(encoding="utf-8")
    privacy_row = f"| 3 | Master Privacy and Data Protection Model | `{rel(PRIVACY_MD)}` | `FOUNDER APPROVED`; format counterpart complete; adoption pending; not locked | Governs privacy and data-protection purpose, minimization, sensitive information, rights, retention, processors, transfers, AI, audit, complaints, and remedies. No implementation or production authority. |"
    reporting_row = f"| 3 | Master Reporting, Analytics, and Business Intelligence Model | `{rel(REPORTING_MD)}` | `FOUNDER APPROVED`; format counterpart complete; adoption pending; not locked | Governs report truth, permissions, provenance, data quality, certification, correction, analytics, research, exports, and public-reporting boundaries. No implementation or production authority. |"
    if privacy_row not in index_text:
        anchor = "| 3 | Master Search, Discovery, Ranking, and Retrieval Model"
        position = index_text.find("\n", index_text.find(anchor))
        index_text = index_text[:position + 1] + privacy_row + "\n" + reporting_row + "\n" + index_text[position + 1:]
    index_text = index_text.replace("C0-023 and C0-035 remain exact-source blocked.", "C0-023 and C0-035 now have exact substantive sources and completed deterministic format counterparts; adoption and lock remain separately gated.")
    write(canon_index, index_text)

    inventory = REPO / "docs/canon/registries/CANON_ARTIFACT_INVENTORY.md"
    inventory_text = inventory.read_text(encoding="utf-8")
    inventory_rows = [
        f"| Master Privacy and Data Protection Model founder-approved Markdown | `{rel(PRIVACY_MD)}` | 2.0 | `{privacy_hash}` | `FOUNDER APPROVED`; counterpart complete; adoption pending; not locked | Status-normalized successor from exact mounted Markdown; generated DOCX is a deterministic parity artifact, not recovered history |",
        f"| Master Reporting, Analytics, and Business Intelligence Model founder-approved DOCX | `{rel(REPORTING_DOCX)}` | 2.0 | `{reporting_hash}` | `FOUNDER APPROVED`; controlling substantive source; adoption pending; not locked | Exact mounted DOCX copied byte-for-byte; Markdown is a deterministic parity artifact, not recovered history |",
    ]
    if inventory_rows[0] not in inventory_text:
        anchor = "| Global Constitutional Companion Artifact Family"
        position = inventory_text.find("\n", inventory_text.find(anchor))
        inventory_text = inventory_text[:position + 1] + "\n".join(inventory_rows) + "\n" + inventory_text[position + 1:]
        write(inventory, inventory_text)

    return [str(path.relative_to(REPO)) for path in [source_register, state_register, row_ledger, delta, *paths.values(), canon_index, inventory]]


def build() -> None:
    if sha256(PRIVACY_SOURCE) != PRIVACY_SOURCE_SHA:
        raise SystemExit("Privacy source hash mismatch")
    if sha256(REPORTING_SOURCE) != REPORTING_SOURCE_SHA:
        raise SystemExit("Reporting source hash mismatch")

    SOURCES.mkdir(parents=True, exist_ok=True)
    APPROVED.mkdir(parents=True, exist_ok=True)
    privacy_source_copy = SOURCES / PRIVACY_SOURCE.name
    reporting_source_copy = SOURCES / REPORTING_SOURCE.name
    shutil.copy2(PRIVACY_SOURCE, privacy_source_copy)
    shutil.copy2(REPORTING_SOURCE, reporting_source_copy)

    privacy_source_text = PRIVACY_SOURCE.read_text(encoding="utf-8")
    privacy_text = privacy_successor(privacy_source_text)
    write(PRIVACY_MD, privacy_text)
    markdown_to_docx(privacy_text, PRIVACY_DOCX, "Master Privacy and Data Protection Model V2.0")

    reporting_source_markdown = docx_to_markdown(REPORTING_SOURCE)
    reporting_text = reporting_successor(reporting_source_markdown)
    write(REPORTING_MD, reporting_text)
    shutil.copy2(REPORTING_SOURCE, REPORTING_DOCX)

    privacy_source_body = "## Authority Boundary" + privacy_source_text.split("## Authority Boundary", 1)[1]
    privacy_successor_body = "## Authority Boundary" + privacy_text.split("## Authority Boundary", 1)[1]
    privacy_substantive_equal = privacy_source_body == privacy_successor_body

    reporting_source_units = docx_units(REPORTING_SOURCE)
    reporting_md_source = reporting_text.split("# Exact Substantive Source Content", 1)[1]
    reporting_derived_units = markdown_units(reporting_md_source)
    reporting_source_compare_text = "\n".join(
        reporting_source_units["paragraphs"]
        + [cell for table in reporting_source_units["tables"] for row in table for cell in row]
    )
    reporting_paragraphs_equal = reporting_source_units["paragraphs"] == reporting_derived_units["paragraphs"]
    reporting_tables_equal = reporting_source_units["tables"] == reporting_derived_units["tables"]

    privacy_docx_data = docx_units(PRIVACY_DOCX)
    privacy_md_data = markdown_units(privacy_text)
    privacy_paragraphs_equal = privacy_md_data["paragraphs"] == privacy_docx_data["paragraphs"]
    privacy_tables_equal = privacy_md_data["tables"] == privacy_docx_data["tables"]

    privacy_metrics = {
        "source_sha256": PRIVACY_SOURCE_SHA,
        "source_size_bytes": PRIVACY_SOURCE.stat().st_size,
        "successor_markdown_sha256": sha256(PRIVACY_MD),
        "successor_docx_sha256": sha256(PRIVACY_DOCX),
        "source_substantive_body_exact_match": privacy_substantive_equal,
        "markdown_paragraphs": len(privacy_md_data["paragraphs"]),
        "docx_paragraphs": len(privacy_docx_data["paragraphs"]),
        "paragraph_sequence_equal": privacy_paragraphs_equal,
        "markdown_headings": len(privacy_md_data["headings"]),
        "docx_headings": len(privacy_docx_data["headings"]),
        "tables": len(privacy_docx_data["tables"]),
        "table_content_equal": privacy_tables_equal,
        "source_requirement_ids": identifiers(privacy_source_body),
        "successor_requirement_ids": identifiers(privacy_successor_body),
        "source_defined_terms": defined_terms(privacy_source_body),
        "successor_defined_terms": defined_terms(privacy_successor_body),
        "source_authority_signals": authority_signals(privacy_source_body),
        "successor_authority_signals": authority_signals(privacy_successor_body),
    }
    reporting_metrics = {
        "source_sha256": REPORTING_SOURCE_SHA,
        "source_size_bytes": REPORTING_SOURCE.stat().st_size,
        "successor_markdown_sha256": sha256(REPORTING_MD),
        "approved_docx_sha256": sha256(REPORTING_DOCX),
        "approved_docx_byte_identical_to_source": sha256(REPORTING_DOCX) == REPORTING_SOURCE_SHA,
        "source_paragraphs": len(reporting_source_units["paragraphs"]),
        "derived_markdown_source_paragraphs": len(reporting_derived_units["paragraphs"]),
        "paragraph_sequence_equal": reporting_paragraphs_equal,
        "source_headings": len(reporting_source_units["headings"]),
        "derived_markdown_headings": len(reporting_derived_units["headings"]),
        "source_tables": len(reporting_source_units["tables"]),
        "derived_markdown_tables": len(reporting_derived_units["tables"]),
        "table_content_equal": reporting_tables_equal,
        "source_requirement_ids": identifiers(reporting_source_compare_text),
        "derived_requirement_ids": identifiers(reporting_md_source),
        "source_defined_terms": defined_terms(reporting_source_compare_text),
        "derived_defined_terms": defined_terms(reporting_md_source),
        "source_authority_signals": authority_signals(reporting_source_compare_text),
        "derived_authority_signals": authority_signals(reporting_md_source),
    }

    required_pass = all([
        privacy_substantive_equal, privacy_paragraphs_equal, privacy_tables_equal,
        reporting_paragraphs_equal, reporting_tables_equal,
        privacy_metrics["source_requirement_ids"] == privacy_metrics["successor_requirement_ids"],
        privacy_metrics["source_defined_terms"] == privacy_metrics["successor_defined_terms"],
        privacy_metrics["source_authority_signals"] == privacy_metrics["successor_authority_signals"],
        reporting_metrics["source_requirement_ids"] == reporting_metrics["derived_requirement_ids"],
        reporting_metrics["source_defined_terms"] == reporting_metrics["derived_defined_terms"],
        reporting_metrics["source_authority_signals"] == reporting_metrics["derived_authority_signals"],
        reporting_metrics["approved_docx_byte_identical_to_source"],
    ])
    if not required_pass:
        write(REVIEW / "PRELIMINARY_PARITY_FAILURE.json", json.dumps({"privacy": privacy_metrics, "reporting": reporting_metrics}, indent=2))
        raise SystemExit("Parity validation failed; see PRELIMINARY_PARITY_FAILURE.json")

    privacy_provenance = f"""# Privacy Format-Counterpart Provenance Record

- C0 row: `C0-023`
- Exact substantive source: `{rel(privacy_source_copy)}`
- Source SHA-256: `{PRIVACY_SOURCE_SHA}`
- Generated Markdown successor: `{rel(PRIVACY_MD)}`
- Generated DOCX counterpart: `{rel(PRIVACY_DOCX)}`
- Evidentiary classification: `{CLASSIFICATION}`
- Historical-recovery claim: `FALSE`
- Substantive source authority: exact mounted Markdown
- Adoption: `NOT YET ADOPTED`
- Lock: `NOT YET LOCKED`
- Operational authority: all false
"""
    reporting_provenance = f"""# Reporting Format-Counterpart Provenance Record

- C0 row: `C0-035`
- Exact substantive source: `{rel(reporting_source_copy)}`
- Source SHA-256: `{REPORTING_SOURCE_SHA}`
- Generated Markdown counterpart: `{rel(REPORTING_MD)}`
- Approved DOCX source alias: `{rel(REPORTING_DOCX)}`
- Evidentiary classification: `{CLASSIFICATION}`
- Historical-recovery claim: `FALSE`
- Substantive source authority: exact mounted DOCX
- Approved DOCX alias byte-identical to source: `TRUE`
- Adoption: `NOT YET ADOPTED`
- Lock: `NOT YET LOCKED`
- Operational authority: all false
"""
    write(REVIEW / "PRIVACY_FORMAT_COUNTERPART_PROVENANCE_RECORD.md", privacy_provenance)
    write(REVIEW / "REPORTING_FORMAT_COUNTERPART_PROVENANCE_RECORD.md", reporting_provenance)

    write(REVIEW / "PRIVACY_MD_DOCX_PARITY_REPORT.md", f"""# Privacy Markdown/DOCX Parity Report

**Result:** `PASS`

- Exact substantive body preserved byte-for-text: `{str(privacy_substantive_equal).upper()}`
- Markdown/DOCX paragraph sequence equal: `{str(privacy_paragraphs_equal).upper()}` ({privacy_metrics['markdown_paragraphs']} / {privacy_metrics['docx_paragraphs']})
- Heading count: `{privacy_metrics['markdown_headings']}` / `{privacy_metrics['docx_headings']}`
- Table count and cells equal: `{str(privacy_tables_equal).upper()}` ({privacy_metrics['tables']})
- Requirement identifiers changed: `0`
- Defined terms changed: `0`
- Authority/prohibition signal deltas: `0`
- Omitted substantive provisions: `0`
- Added substantive provisions: `0`
- Unexplained semantic deltas: `0`

The generated DOCX is a deterministic format counterpart, not a recovered historical source.
""")
    write(REVIEW / "REPORTING_DOCX_MD_PARITY_REPORT.md", f"""# Reporting DOCX/Markdown Parity Report

**Result:** `PASS`

- Source/derived paragraph sequence equal: `{str(reporting_paragraphs_equal).upper()}` ({reporting_metrics['source_paragraphs']} / {reporting_metrics['derived_markdown_source_paragraphs']})
- Heading count: `{reporting_metrics['source_headings']}` / `{reporting_metrics['derived_markdown_headings']}`
- Table count and cells equal: `{str(reporting_tables_equal).upper()}` ({reporting_metrics['source_tables']} / {reporting_metrics['derived_markdown_tables']})
- Approved DOCX alias byte-identical to exact source: `{str(reporting_metrics['approved_docx_byte_identical_to_source']).upper()}`
- Requirement identifiers changed: `0`
- Defined terms changed: `0`
- Authority/prohibition signal deltas: `0`
- Omitted substantive provisions: `0`
- Added substantive provisions: `0`
- Unexplained semantic deltas: `0`

The generated Markdown is a deterministic format counterpart, not a recovered historical source.
""")

    completion_report = f"""# Privacy and Reporting Format-Counterpart Completion Report

## Disposition

`{FINAL_DISPOSITION}`

Both exact substantive source authorities passed hash verification. The Privacy founder-approved Markdown and DOCX pair and Reporting founder-approved DOCX and Markdown pair are complete. Generated formats are classified as `{CLASSIFICATION}` and are not represented as recovered historical sources.

## Results

| Row | Exact source | Exact source SHA-256 | Generated counterpart | Parity |
|---|---|---|---|---|
| C0-023 | `{rel(privacy_source_copy)}` | `{PRIVACY_SOURCE_SHA}` | `{rel(PRIVACY_DOCX)}` | PASS |
| C0-035 | `{rel(reporting_source_copy)}` | `{REPORTING_SOURCE_SHA}` | `{rel(REPORTING_MD)}` | PASS |

Omitted substantive provisions: `0`. Added substantive provisions: `0`. Changed requirement identifiers: `0`. Changed authority boundaries: `0`. Unexplained semantic deltas: `0`.

## Lifecycle

Exact substantive source is available, founder approval is confirmed, deterministic format counterpart completion is complete, and the source-reconstruction blocker is resolved for C0-023 and C0-035. Both rows remain `adoption_and_lock_required`; adoption and lock require separate founder authorization.

## Authority Boundary

Adoption, lock, Governance V1.0 adoption/lock, implementation, migration, runtime, production, provider activation, deployment, customer-data processing, public launch, security certification, and public trust claims remain `FALSE`.
"""
    write(REVIEW / "PRIVACY_REPORTING_FORMAT_COUNTERPART_COMPLETION_REPORT.md", completion_report)

    companion_paths = update_companions(sha256(PRIVACY_MD), sha256(REPORTING_DOCX))

    artifacts = [privacy_source_copy, reporting_source_copy, PRIVACY_MD, PRIVACY_DOCX, REPORTING_MD, REPORTING_DOCX]
    manifest = {
        "disposition": FINAL_DISPOSITION,
        "classification": CLASSIFICATION,
        "privacy": privacy_metrics,
        "reporting": reporting_metrics,
        "artifacts": [{"path": rel(path), "size_bytes": path.stat().st_size, "sha256": sha256(path)} for path in artifacts],
        "companion_artifacts_updated": companion_paths,
        "visual_validation": "PENDING_SEPARATE_RENDER_GATE",
        "authority": AUTHORITY,
    }
    write(REVIEW / "PRIVACY_REPORTING_FORMAT_COUNTERPART_MANIFEST.json", json.dumps(manifest, indent=2))
    print(json.dumps({"disposition": FINAL_DISPOSITION, "privacy": privacy_metrics, "reporting": reporting_metrics}, indent=2))


if __name__ == "__main__":
    build()
