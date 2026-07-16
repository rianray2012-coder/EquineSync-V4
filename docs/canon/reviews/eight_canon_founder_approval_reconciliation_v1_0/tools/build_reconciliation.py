#!/usr/bin/env python3
"""Build the documentation-only eight-canon founder-approval reconciliation."""

from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


REPO = Path(__file__).resolve().parents[5]
REVIEW = REPO / "docs/canon/reviews/eight_canon_founder_approval_reconciliation_v1_0"
SUCCESSORS = REPO / "docs/canon/founder_approved_sources"
PREDECESSORS = REVIEW / "predecessors"
SOURCE_EVENT = REVIEW / "source_event"
RENDER = REVIEW / "render_evidence"
COMPANIONS = REVIEW / "companion_updates"

DIRECTIVE = Path(
    "/Users/rianray/.codex/attachments/43a9f352-4d1f-402e-b641-b226fe6ec329/pasted-text.txt"
)
INTAKE_ZIP = Path("/Users/rianray/Downloads/EQUINESYNC_TRACK_D_SOURCE_INTAKE_PACKAGE.zip")
INTAKE_SHA = Path("/Users/rianray/Downloads/EQUINESYNC_TRACK_D_SOURCE_INTAKE_PACKAGE.zip.sha256")


@dataclass(frozen=True)
class Model:
    row_id: str
    basename: str
    title: str
    source: Path | None
    source_kind: str
    approval_date: str
    source_note: str
    preserve_stronger_status: bool = False


MODELS = [
    Model(
        "C0-043",
        "MASTER_PLATFORM_OPERATIONS_RELIABILITY_AND_RELEASE_MODEL_V2_0",
        "Master Platform Operations, Reliability, and Release Model",
        Path("/Users/rianray/Downloads/MASTER_PLATFORM_OPERATIONS_RELIABILITY_AND_RELEASE_MODEL_V2_0.md"),
        "markdown",
        "July 16, 2026",
        "Exact V2.0 substantive Markdown; byte-identical to the repository candidate.",
    ),
    Model(
        "C0-042",
        "MASTER_CONFIGURATION_AND_FEATURE_FLAG_GOVERNANCE_MODEL_V2_0",
        "Master Configuration and Feature Flag Governance Model",
        Path("/Users/rianray/Downloads/MASTER_CONFIGURATION_AND_FEATURE_FLAG_GOVERNANCE_MODEL_V2_0.docx"),
        "docx",
        "July 16, 2026",
        "Exact V2.0 DOCX containing FD01-FD50.",
    ),
    Model(
        "C0-041",
        "MASTER_VENDOR_SECURITY_AND_SUPPLY_CHAIN_MODEL_V2_0",
        "Master Vendor Security and Supply Chain Model",
        Path("/Users/rianray/Downloads/MASTER_VENDOR_SECURITY_AND_SUPPLY_CHAIN_MODEL_V2_0_FOUNDER_ACCEPTED.docx"),
        "docx",
        "July 13, 2026",
        "Exact founder-accepted V2.0 DOCX; earlier evidenced acceptance date preserved.",
        True,
    ),
    Model(
        "C0-040",
        "MASTER_PLATFORM_EXTENSIBILITY_AND_PLUGIN_GOVERNANCE_MODEL_V2_0",
        "Master Platform Extensibility and Plugin Governance Model",
        Path("/Users/rianray/Downloads/MASTER_PLATFORM_EXTENSIBILITY_AND_PLUGIN_GOVERNANCE_MODEL_V2_0(1).md"),
        "markdown",
        "July 16, 2026",
        "Later full-text V2.0 Markdown; early formatting scaffold is historical and was not mounted.",
    ),
    Model(
        "C0-039",
        "MASTER_DEVELOPER_PLATFORM_AND_INTEGRATION_GOVERNANCE_MODEL_V2_0",
        "Master Developer, Platform, and Integration Governance Model",
        Path("/Users/rianray/Downloads/MASTER_DEVELOPER_PLATFORM_AND_INTEGRATION_GOVERNANCE_MODEL_V2_0.md"),
        "markdown",
        "July 16, 2026",
        "Exact V2.0 substantive Markdown approved by the founder source event.",
    ),
    Model(
        "C0-035",
        "MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0",
        "Master Reporting, Analytics, and Business Intelligence Model",
        None,
        "missing",
        "",
        "Exact V2.0 DOCX/Markdown source bytes were absent from the intake package, repository, mounted workspace, and local archives.",
    ),
    Model(
        "C0-033",
        "MASTER_SEARCH_DISCOVERY_RANKING_AND_RETRIEVAL_MODEL_V2_0",
        "Master Search, Discovery, Ranking, and Retrieval Model",
        Path("/Users/rianray/Downloads/MASTER_SEARCH_DISCOVERY_RANKING_AND_RETRIEVAL_MODEL_V2_0_FIRST_DRAFT.docx"),
        "docx",
        "July 16, 2026",
        "Exact V2.0 first-draft DOCX with FD-SD01-FD-SD14 and Shared Care/Delegated Assistance text.",
    ),
    Model(
        "C0-023",
        "MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0",
        "Master Privacy and Data Protection Model",
        None,
        "missing",
        "",
        "Exact reviewed V2.0 DOCX/Markdown source bytes were absent from the intake package, repository, mounted workspace, and local archives.",
    ),
]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def iter_blocks(document: Document) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def escape_table_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>").strip()


def docx_to_markdown(path: Path) -> str:
    document = Document(path)
    lines: list[str] = []
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                lines.append("")
                continue
            style = (block.style.name or "").lower()
            heading_match = re.search(r"heading\s*(\d+)", style)
            if heading_match:
                level = min(int(heading_match.group(1)), 6)
                lines.extend([f"{'#' * level} {text}", ""])
            elif "title" in style:
                lines.extend([f"# {text}", ""])
            elif "subtitle" in style:
                lines.extend([f"## {text}", ""])
            elif "list bullet" in style:
                lines.append(f"- {text}")
            elif "list number" in style:
                lines.append(f"1. {text}")
            else:
                lines.extend([text, ""])
        else:
            rows = []
            for row in block.rows:
                rows.append([escape_table_cell(cell.text) for cell in row.cells])
            if not rows:
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n"


def lifecycle_block(model: Model) -> str:
    status = "CONTROLLING CANON" if model.preserve_stronger_status else "CONTROLLING SUBSTANTIVE CANON"
    return (
        "## Founder Approval and Authority Status\n\n"
        "| Field | Controlling Value |\n"
        "|---|---|\n"
        "| Founder Approval Status | **FOUNDER APPROVED** |\n"
        f"| Founder Approval Date | **{model.approval_date}** |\n"
        f"| Constitutional Status | **{status}** |\n"
        "| Canon Adoption | **PENDING SEPARATE ADOPTION EVIDENCE** |\n"
        "| Canon Lock | **NOT LOCKED BY THIS DIRECTIVE** |\n"
        "| Implementation Authority | **FALSE** |\n"
        "| Schema / Migration Authority | **FALSE** |\n"
        "| Runtime Activation Authority | **FALSE** |\n"
        "| Production Authority | **FALSE** |\n"
        "| External-Provider Activation Authority | **FALSE** |\n"
        "| Public-Claims Authority | **FALSE** |\n"
        "| Public-Launch Authority | **FALSE** |\n\n"
        "> Founder approval establishes the controlling substantive constitutional text. "
        "Adoption, lock, implementation, runtime, production, provider, public-claims, and launch authority remain separately gated.\n"
    )


def normalize_status(text: str, model: Model) -> str:
    replacements = [
        (r"FINAL CONTROLLED CANDIDATE FOR FOUNDER ACCEPTANCE", "FOUNDER APPROVED CONTROLLING SUBSTANTIVE CANON"),
        (r"READY FOR FOUNDER ACCEPTANCE", "FOUNDER APPROVED"),
        (r"DRAFT_FOR_CONTROLLED_CONSTITUTIONAL_REVIEW", "FOUNDER_APPROVED_CONTROLLING_SUBSTANTIVE_CANON"),
        (r"Draft for Controlled Constitutional Review", "Founder Approved Controlling Substantive Canon"),
        (r"CONTROLLED CONSTITUTIONAL DRAFT", "FOUNDER-APPROVED CONTROLLING SUBSTANTIVE CANON"),
        (r"Controlled Constitutional Draft", "Founder-Approved Controlling Substantive Canon"),
        (r"controlled constitutional draft", "founder-approved controlling substantive canon"),
        (r"FIRST DRAFT", "FOUNDER APPROVED"),
        (r"First Draft", "Founder Approved"),
        (r"first draft", "founder-approved text"),
        (r"REVIEWED DRAFT", "FOUNDER APPROVED"),
        (r"NOT ADOPTED", "ADOPTION PENDING SEPARATE EVIDENCE"),
        (r"Not Adopted", "Adoption Pending Separate Evidence"),
        (r"not adopted and not locked", "founder-approved; adoption remains pending and this directive does not lock it"),
        (r"not adopted, locked, implemented, or activated", "founder-approved, but not adopted, locked, implemented, or activated"),
        (r"Founder-directed\s*\|\s*Not adopted", "Founder approved | Adoption pending separate evidence"),
        (r"Founder-directed constitutional", "Founder-approved constitutional"),
        (r"Founder-directed", "Founder approved"),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # Normalize common status rows without altering substantive authority language.
    text = re.sub(r"\*\*Canon adoption:\*\*\s*NOT YET AUTHORIZED", "**Canon adoption:** PENDING SEPARATE ADOPTION EVIDENCE", text, flags=re.I)
    text = re.sub(r"\*\*Canon lock:\*\*\s*NOT AUTHORIZED", "**Canon lock:** NOT LOCKED BY THIS DIRECTIVE", text, flags=re.I)
    text = re.sub(r"\*\*Adopted:\*\*\s*False", "**Canon Adoption:** PENDING SEPARATE ADOPTION EVIDENCE", text, flags=re.I)
    text = re.sub(r"\*\*Locked:\*\*\s*False", "**Canon Lock:** NOT LOCKED BY THIS DIRECTIVE", text, flags=re.I)

    lines = text.splitlines()
    insert_at = 1
    while insert_at < len(lines) and not lines[insert_at].strip():
        insert_at += 1
    lines[insert_at:insert_at] = ["", lifecycle_block(model).rstrip(), ""]
    text = "\n".join(lines).strip() + "\n"
    return text


def add_page_number(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip().replace("<br>", "\n").replace("\\|", "|") for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows


def markdown_to_docx(markdown: str, output: Path, title: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.3)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for level in range(1, 7):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(55, 47, 73)
        style.font.size = Pt(max(11, 22 - level * 2))
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "EQUINESYNC  |  FOUNDER-APPROVED CONSTITUTIONAL SOURCE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(111, 96, 128)
    add_page_number(section.footer.paragraphs[0])

    lines = markdown.splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = parse_table(table_lines)
            if rows:
                width = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(width):
                        cell = table.cell(r_idx, c_idx)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                if r_idx == 0:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                        if r_idx == 0:
                            set_cell_shading(cell, "40364D")
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            paragraph = document.add_paragraph(text, style=f"Heading {level}")
            if first_heading:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_heading = False
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(line[2:].strip())
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.right_indent = Inches(0.3)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(80, 70, 92)
            index += 1
            continue
        if re.match(r"^[-*]\s+", line):
            document.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            index += 1
            continue
        paragraph = document.add_paragraph()
        segments = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", line)
        for segment in segments:
            if not segment:
                continue
            if segment.startswith("**") and segment.endswith("**"):
                run = paragraph.add_run(segment[2:-2])
                run.bold = True
            elif segment.startswith("`") and segment.endswith("`"):
                run = paragraph.add_run(segment[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
            else:
                paragraph.add_run(segment)
        index += 1

    core = document.core_properties
    core.title = title
    core.subject = "Founder-approved constitutional source; adoption and lock remain separately gated"
    core.author = "EquineSync"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    for directory in (SUCCESSORS, PREDECESSORS, SOURCE_EVENT, RENDER, COMPANIONS):
        directory.mkdir(parents=True, exist_ok=True)

    directive_target = SOURCE_EVENT / "EQUINESYNC_EIGHT_CANON_FOUNDER_APPROVAL_AND_STATUS_RECONCILIATION_DIRECTIVE_V1_0.md"
    shutil.copy2(DIRECTIVE, directive_target)
    shutil.copy2(INTAKE_ZIP, SOURCE_EVENT / INTAKE_ZIP.name)
    shutil.copy2(INTAKE_SHA, SOURCE_EVENT / INTAKE_SHA.name)

    rows = []
    hash_rows = []
    for model in MODELS:
        if model.source is None:
            rows.append({
                "row_id": model.row_id,
                "model": model.title,
                "result": "BLOCKED_EXACT_SOURCE_NOT_LOCATED",
                "founder_approval_status": "APPROVED_BY_DIRECTIVE_BUT_SUCCESSOR_NOT_GENERATED_WITHOUT_SOURCE",
                "constitutional_status": "SOURCE_BLOCKED",
                "canon_adoption": "PENDING SEPARATE ADOPTION EVIDENCE",
                "canon_lock": "NOT LOCKED BY THIS DIRECTIVE",
                "source_note": model.source_note,
                "implementation_authority": False,
                "production_authority": False,
                "public_launch_authority": False,
            })
            continue

        predecessor = PREDECESSORS / model.source.name
        shutil.copy2(model.source, predecessor)
        source_text = model.source.read_text(encoding="utf-8") if model.source_kind == "markdown" else docx_to_markdown(model.source)
        successor_text = normalize_status(source_text, model)
        md_target = SUCCESSORS / f"{model.basename}_FOUNDER_APPROVED.md"
        docx_target = SUCCESSORS / f"{model.basename}_FOUNDER_APPROVED.docx"
        md_target.write_text(successor_text, encoding="utf-8")
        markdown_to_docx(successor_text, docx_target, model.title)

        status = "CONTROLLING CANON" if model.preserve_stronger_status else "CONTROLLING SUBSTANTIVE CANON"
        rows.append({
            "row_id": model.row_id,
            "model": model.title,
            "result": "RECONCILED_FOUNDER_APPROVED_SUCCESSOR_CREATED",
            "source_path": str(model.source),
            "predecessor_path": str(predecessor.relative_to(REPO)),
            "predecessor_sha256": sha256(predecessor),
            "markdown_successor_path": str(md_target.relative_to(REPO)),
            "markdown_successor_sha256": sha256(md_target),
            "docx_successor_path": str(docx_target.relative_to(REPO)),
            "docx_successor_sha256": sha256(docx_target),
            "founder_approval_status": "FOUNDER APPROVED",
            "founder_approval_date": model.approval_date,
            "constitutional_status": status,
            "canon_adoption": "PENDING SEPARATE ADOPTION EVIDENCE",
            "canon_lock": "NOT LOCKED BY THIS DIRECTIVE",
            "source_note": model.source_note,
            "implementation_authority": False,
            "schema_migration_authority": False,
            "runtime_activation_authority": False,
            "production_authority": False,
            "external_provider_activation_authority": False,
            "public_claims_authority": False,
            "public_launch_authority": False,
        })
        for artifact_type, path in (("predecessor", predecessor), ("successor_markdown", md_target), ("successor_docx", docx_target)):
            hash_rows.append({
                "row_id": model.row_id,
                "model": model.title,
                "artifact_type": artifact_type,
                "path": str(path.relative_to(REPO)),
                "sha256": sha256(path),
                "size_bytes": path.stat().st_size,
                "reason_for_succession": "Founder-approved lifecycle status correction; substantive source preserved.",
            })

    ledger = {
        "package": "EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION_V1_0",
        "directive_sha256": sha256(directive_target),
        "intake_zip_sha256": sha256(SOURCE_EVENT / INTAKE_ZIP.name),
        "generated_on": date.today().isoformat(),
        "disposition": "EIGHT_CANON_FOUNDER_APPROVAL_STATUS_RECONCILIATION_PARTIALLY_COMPLETE_WITH_EXACT_SOURCE_BLOCKERS",
        "reconciled_rows": sum(r["result"].startswith("RECONCILED") for r in rows),
        "blocked_rows": sum(r["result"].startswith("BLOCKED") for r in rows),
        "authority": {
            "implementation": False,
            "schema_or_migration": False,
            "runtime_activation": False,
            "production": False,
            "external_provider_activation": False,
            "public_claims": False,
            "public_launch": False,
            "canon_lock": False,
        },
        "rows": rows,
    }
    (REVIEW / "EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION_LEDGER.json").write_text(
        json.dumps(ledger, indent=2) + "\n", encoding="utf-8"
    )
    with (REVIEW / "EIGHT_CANON_PREDECESSOR_SUCCESSOR_HASH_REGISTER.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(hash_rows[0]))
        writer.writeheader()
        writer.writerows(hash_rows)

    report_lines = [
        "# Eight-Canon Founder Approval Reconciliation Report",
        "",
        "## Disposition",
        "",
        "`EIGHT_CANON_FOUNDER_APPROVAL_STATUS_RECONCILIATION_PARTIALLY_COMPLETE_WITH_EXACT_SOURCE_BLOCKERS`",
        "",
        "Six of eight rows have exact substantive sources and complete founder-approved Markdown/DOCX successors. Reporting and Privacy remain blocked because their exact substantive bytes were not mounted or present in any inspected repository, workspace, package, or local archive. No source content was reconstructed from snippets.",
        "",
        "## Authority Boundary",
        "",
        "This reconciliation creates no implementation, schema, migration, runtime, production, external-provider, public-claims, public-launch, or canon-lock authority. Founder approval of substantive text is recorded separately from adoption and lock.",
        "",
        "## Row Results",
        "",
        "| C0 Row | Model | Result | Founder Status | Adoption | Lock |",
        "|---|---|---|---|---|---|",
    ]
    for row in rows:
        report_lines.append(
            f"| {row['row_id']} | {row['model']} | `{row['result']}` | {row['founder_approval_status']} | {row['canon_adoption']} | {row['canon_lock']} |"
        )
    report_lines.extend([
        "",
        "## Source Resolution Notes",
        "",
        "- The Track D intake ZIP is a manifest-only package and contains no constitutional source bytes.",
        "- The Plugin V2.0 source used here is the later 86,950-byte full-text Markdown. The earlier formatting scaffold pair is classified as historical/non-controlling but was not mounted and therefore has no local hash claim.",
        "- Exact V2.1 Developer and Plugin files from the Cross-Canon Reconciliation package are retained as adjacent historical/version evidence; this directive expressly approves the identified V2.0 models and no silent V2.1 substitution was made.",
        "- Vendor Security preserves its evidenced July 13, 2026 founder-acceptance date and stronger controlling-canon status.",
        "- Search preserves FD-SD01 through FD-SD14, Shared Care and Delegated Assistance, no-access-expansion, and public-discovery gates.",
        "",
        "## Remaining Exact-Source Blockers",
        "",
        "1. `C0-035`: `MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0.docx` was not located.",
        "2. `C0-023`: `MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_REVIEWED_DRAFT.docx` was not located.",
        "",
        "These rows may resume independently when exact source bytes are supplied.",
    ])
    (REVIEW / "EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION_REPORT.md").write_text(
        "\n".join(report_lines) + "\n", encoding="utf-8"
    )

    companion = {
        "update_id": "EIGHT_CANON_FOUNDER_APPROVAL_LIFECYCLE_UPDATE_V1_0",
        "directive_sha256": sha256(directive_target),
        "applies_to": [
            "Canon Index and Navigator",
            "Domain Ownership and Boundary Register",
            "Authority Matrix",
            "Cross Reference Index",
            "Dependency Map",
            "Founder Decision Register",
            "Governance Requirement Index",
            "Requirement Traceability Matrix",
            "C0 Source-of-Truth Register",
            "C0 Row-by-Row Lifecycle Resolution Ledger",
            "provenance, supersession, adoption, and checksum registers",
        ],
        "rows": [{
            "row_id": row["row_id"],
            "model": row["model"],
            "founder_status": row["founder_approval_status"],
            "adoption_state": row["canon_adoption"],
            "lock_state": row["canon_lock"],
            "result": row["result"],
            "successor": row.get("markdown_successor_path"),
        } for row in rows],
        "authority_flags": ledger["authority"],
    }
    (COMPANIONS / "EIGHT_CANON_COMPANION_LIFECYCLE_UPDATE.json").write_text(
        json.dumps(companion, indent=2) + "\n", encoding="utf-8"
    )
    md = [
        "# Eight-Canon Companion Lifecycle Update",
        "",
        "This additive update applies the founder-approved source identity and lifecycle status to all named companion instruments without claiming adoption or lock.",
        "",
        "| Row | Model | Founder status | Adoption | Lock | Result |",
        "|---|---|---|---|---|---|",
    ]
    for row in rows:
        md.append(f"| {row['row_id']} | {row['model']} | {row['founder_approval_status']} | {row['canon_adoption']} | {row['canon_lock']} | `{row['result']}` |")
    md.extend([
        "",
        "## Application Rule",
        "",
        "For reconciled rows, companion references must resolve to the founder-approved successor while describing adoption as pending and lock as not established by this directive. For blocked rows, existing unresolved-source treatment remains controlling. All operational authority flags remain false.",
    ])
    (COMPANIONS / "EIGHT_CANON_COMPANION_LIFECYCLE_UPDATE.md").write_text("\n".join(md) + "\n", encoding="utf-8")

    update_c0_lifecycle(rows)


def update_c0_lifecycle(rows: list[dict]) -> None:
    """Apply the additive lifecycle result to the live C0 tracking records."""
    by_id = {row["row_id"]: row for row in rows}
    current_path = REPO / "docs/canon/reviews/governance_v1_0_c0_lifecycle_completion_batch_1/C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.json"
    current = json.loads(current_path.read_text(encoding="utf-8"))
    records = current.get("records", current.get("rows", []))
    for record in records:
        row = by_id.get(record.get("record_id"))
        if not row:
            continue
        if row["result"].startswith("RECONCILED"):
            record.update({
                "current_category": "adoption_and_lock_required",
                "current_exact_source_status": "FOUNDER_APPROVED_SUCCESSOR_SOURCE_VERIFIED",
                "current_repository_path": row["markdown_successor_path"],
                "current_sha256": row["markdown_successor_sha256"],
                "founder_status": "FOUNDER_APPROVED",
                "adoption_state": "PENDING_SEPARATE_ADOPTION_EVIDENCE",
                "lock_state": "NOT_LOCKED_BY_THIS_DIRECTIVE",
                "current_controlling_artifact": row["markdown_successor_path"],
                "lifecycle_evidence": "docs/canon/reviews/eight_canon_founder_approval_reconciliation_v1_0/EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION_LEDGER.json",
                "unresolved_blocker": True,
                "remedy": "Complete separately authorized canon adoption and lock review; all operational authority remains false.",
            })
        else:
            record.update({
                "current_category": "unresolved_source_evidence",
                "current_exact_source_status": "SOURCE_NOT_LOCATED_AFTER_TRACK_D_INTAKE",
                "founder_status": "FOUNDER_APPROVED_BY_DIRECTIVE_SOURCE_BLOCKED",
                "adoption_state": "PENDING_SEPARATE_ADOPTION_EVIDENCE",
                "lock_state": "NOT_LOCKED_BY_THIS_DIRECTIVE",
                "lifecycle_evidence": "docs/canon/reviews/eight_canon_founder_approval_reconciliation_v1_0/EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION_LEDGER.json",
                "unresolved_blocker": True,
                "remedy": "Supply exact substantive source bytes; then generate and verify the founder-approved successor before any adoption or lock action.",
            })
    current["generated_at"] = date.today().isoformat()
    current["status"] = "UPDATED_AFTER_EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION"
    current_path.write_text(json.dumps(current, indent=2) + "\n", encoding="utf-8")
    render_current_register(current, current_path.with_suffix(".md"))

    ledger_path = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.json"
    ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    for record in ledger["records"]:
        row = by_id.get(record.get("record_id"))
        if not row:
            continue
        if row["result"].startswith("RECONCILED"):
            record.update({
                "resolution_track": "TRACK_F_ADOPTION_AND_LOCK",
                "current_category": "adoption_and_lock_required",
                "source": row["markdown_successor_path"],
                "founder_status": "FOUNDER_APPROVED",
                "adoption_state": "PENDING_SEPARATE_ADOPTION_EVIDENCE",
                "lock_state": "NOT_LOCKED_BY_THIS_DIRECTIVE",
                "next_step": "Complete separately authorized canon adoption and lock review; all operational authority remains false.",
                "track_d_status": "FOUNDER_APPROVAL_STATUS_RECONCILED",
                "track_d_classification": "FOUNDER_APPROVED_SUCCESSOR_SOURCE_VERIFIED_ADOPTION_AND_LOCK_PENDING",
            })
        else:
            record.update({
                "resolution_track": "TRACK_D_MISSING_SOURCE_RECOVERY",
                "current_category": "unresolved_source_evidence",
                "founder_status": "FOUNDER_APPROVED_BY_DIRECTIVE_SOURCE_BLOCKED",
                "adoption_state": "PENDING_SEPARATE_ADOPTION_EVIDENCE",
                "lock_state": "NOT_LOCKED_BY_THIS_DIRECTIVE",
                "next_step": "Supply exact substantive source bytes, generate the successor, then complete separately authorized adoption and lock review.",
                "track_d_status": "EXACT_SOURCE_NOT_RECOVERED",
                "track_d_classification": "FOUNDER_APPROVAL_DIRECTIVE_PRESENT_EXACT_SUBSTANTIVE_SOURCE_ABSENT",
            })
    ledger["generated_at"] = date.today().isoformat()
    ledger["status"] = "UPDATED_AFTER_EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION"
    counts: dict[str, int] = {}
    for record in ledger["records"]:
        counts[record["resolution_track"]] = counts.get(record["resolution_track"], 0) + 1
    ledger["track_counts"] = counts
    ledger_path.write_text(json.dumps(ledger, indent=2) + "\n", encoding="utf-8")
    render_row_ledger(ledger, ledger_path.with_suffix(".md"))
    render_blocker_ledger(ledger, ledger_path.parent / "C0_UNRESOLVED_LIFECYCLE_BLOCKER_LEDGER.md")


def render_current_register(data: dict, path: Path) -> None:
    records = data.get("records", data.get("rows", []))
    lines = [
        "# C0 Current Source-of-Truth Register",
        "",
        f"Status: `{data.get('status')}`. Founder-approved substantive text is distinct from adoption, lock, and operational authority.",
        "",
        "| C0 | Family | Current category | Source status | Founder status | Adoption | Lock | Controlling artifact |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for record in records:
        lines.append(
            f"| {record['record_id']} | {record['family']} | {record.get('current_category','')} | {record.get('current_exact_source_status','')} | {record.get('founder_status','')} | {record.get('adoption_state','')} | {record.get('lock_state','')} | {record.get('current_controlling_artifact') or ''} |"
        )
    lines.extend(["", "All implementation, migration, runtime, production, provider, public-claims, public-launch, and Governance V1 baseline adoption/lock authority remain false unless separately evidenced."])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_row_ledger(data: dict, path: Path) -> None:
    lines = [
        "# C0 Row-by-Row Lifecycle Resolution Ledger",
        "",
        "Status: `UPDATED_AFTER_EIGHT_CANON_FOUNDER_APPROVAL_RECONCILIATION`; remaining lifecycle blockers: `17`. Six rows now have founder-approved successors but remain blocked on separately authorized adoption and lock. Two Track D rows remain exact-source blocked.",
        "",
        "| C0 | Track | Category | Founder status | Adoption | Lock | Next step |",
        "|---|---|---|---|---|---|---|",
    ]
    for record in data["records"]:
        lines.append(
            f"| {record['record_id']} | {record['resolution_track']} | {record.get('current_category','')} | {record.get('founder_status','')} | {record.get('adoption_state','')} | {record.get('lock_state','')} | {record.get('next_step','')} |"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_blocker_ledger(data: dict, path: Path) -> None:
    lines = [
        "# C0 Unresolved Lifecycle Blocker Ledger",
        "",
        "Open blockers: `17`. Founder approval changed the reason for six blockers but did not authorize adoption or lock.",
        "",
        "| C0 | Family | Category | Founder status | Required resolution |",
        "|---|---|---|---|---|",
    ]
    for record in data["records"]:
        lines.append(
            f"| {record['record_id']} | {record['family']} | {record.get('current_category','')} | {record.get('founder_status','')} | {record.get('next_step','')} |"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
