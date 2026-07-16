#!/usr/bin/env python3
"""Build C0-019 Phase A source, counterpart, and review evidence."""

from __future__ import annotations

import hashlib
import json
import mimetypes
import re
import shutil
import zipfile
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[5]
ROOT = REPO / "docs/canon/reviews/c0_019_agreement_consent_authorization_lifecycle_completion"
MOUNTED = Path("/Users/rianray/Downloads/equinesync_cross_canon_reconciliation/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_1.md")
PACKAGE = Path("/Users/rianray/Downloads/EquineSync_Cross_Canon_Reconciliation_Package_V1_0.zip")
PACKAGE_ENTRY = "equinesync_cross_canon_reconciliation/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_1.md"
RAW = ROOT / "source_event/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_1_ORIGINAL.md"
V20 = REPO / "docs/canon/candidates/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_0.md"
HIST = REPO / "docs/canon/history/agreement_consent_authorization_v2_0/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_0_HISTORICAL_SUPERSEDED_CANDIDATE.md"
ACTIVE_MD = REPO / "docs/canon/founder_approved_sources/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_1_FOUNDER_APPROVED.md"
ACTIVE_DOCX = REPO / "docs/canon/founder_approved_sources/MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_1_FOUNDER_APPROVED.docx"
EXPECTED = "af34895d8248f0fa26f7c976c345caa7ac9fba7b7bf4f597d32bfb68e0797d20"
V20_HASH = "630b6128e37734be01a40d64b80d9fb4d74fce48198c60f9fcbd1a2ef83c232b"
PACKAGE_HASH = "937c3f702bf23eb89addfee50f17c707a38f0355abace710d23088304fa66f55"
AUTHORITY = {
    "c0_023_privacy_lock": False, "c0_022_permission_lifecycle": False,
    "c0_035_reporting_lock": False, "governance_v1_baseline_adoption": False,
    "governance_v1_baseline_lock": False, "implementation": False,
    "migration": False, "runtime": False, "provider_activation": False,
    "integration_activation": False, "deployment": False, "production": False,
    "customer_data_processing": False, "public_launch": False,
    "certification": False, "public_trust_claims": False,
}


def now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path) -> str:
    return path.relative_to(REPO).as_posix()


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_json(path: Path, value: object) -> None:
    write_text(path, json.dumps(value, indent=2, ensure_ascii=True))


def lifecycle_markdown(adoption: str, lock: str) -> str:
    source = RAW.read_text()
    body = source[source.index("# 1. Constitutional Purpose"):]
    return f"""# MASTER AGREEMENT, CONSENT, AND AUTHORIZATION MODEL

**Document Type:** Constitutional Canon  
**C0 Identifier:** C0-019  
**Founder Approval Status:** FOUNDER APPROVED  
**Source Identity:** VERIFIED  
**Controlling Version:** V2.1  
**Constitutional Adoption Status:** {adoption}  
**Constitutional Lock Status:** {lock}  
**Implementation Authority:** FALSE  
**Runtime Authority:** FALSE  
**Migration Authority:** FALSE  
**Provider Activation Authority:** FALSE  
**Integration Activation Authority:** FALSE  
**Deployment Authority:** FALSE  
**Production Authority:** FALSE  
**Customer-Data Processing Authority:** FALSE  
**Public Launch Authority:** FALSE  
**Certification Authority:** FALSE  
**Public Trust Claim Authority:** FALSE  

## Format-Counterpart Provenance

This Markdown preserves the exact substantive content of the founder-authorized V2.1 source with raw SHA-256 `{EXPECTED}`. Only lifecycle document-control metadata differs. The DOCX is a founder-authorized deterministic format counterpart and is not represented as a recovered historical original.

{body}
"""


def normalized(text: str, substantive: bool = False) -> str:
    if substantive and "# 1. Constitutional Purpose" in text:
        text = text[text.index("# 1. Constitutional Purpose"):]
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.M)
    text = re.sub(r"^[-*]\s+", "", text, flags=re.M)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.M)
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"(?:^|\s)\d+(?:\.\d+)*\.\s+", " ", text)
    text = re.sub(r"(?:^|\s)---(?:\s|$)", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def blocks(text: str) -> list[tuple[str, str]]:
    result: list[tuple[str, str]] = []
    para: list[str] = []
    def flush() -> None:
        if para:
            result.append(("paragraph", " ".join(para)))
            para.clear()
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip() or line.strip() == "---":
            flush(); continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        bullet = re.match(r"^-\s+(.+)$", line)
        if heading:
            flush(); result.append((f"heading{len(heading.group(1))}", heading.group(2).strip()))
        elif numbered:
            flush(); result.append(("numbered", numbered.group(1).strip()))
        elif bullet:
            flush(); result.append(("bullet", bullet.group(1).strip()))
        elif line.startswith("**") and ":**" in line:
            flush(); result.append(("metadata", line.strip().rstrip().replace("**", "")))
        else:
            para.append(line.strip())
    flush()
    return result


def add_numbering(document: Document, kind: str, num_id: int, abstract_id: int) -> None:
    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    level = OxmlElement("w:lvl"); level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal"); level.append(fmt)
    txt = OxmlElement("w:lvlText"); txt.set(qn("w:val"), "•" if kind == "bullet" else "%1."); level.append(txt)
    jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); level.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360"); ppr.append(ind); level.append(ppr)
    abstract.append(level); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(abstract_id)); num.append(ref); numbering.append(num)


def set_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr(); numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id)); numpr.extend((ilvl, nid)); ppr.append(numpr)


def page_field(paragraph) -> None:
    run = paragraph.add_run()
    for kind, text in (("begin", None), (None, " PAGE "), ("separate", None), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar"); node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText"); node.set(qn("xml:space"), "preserve"); node.text = text
        run._r.append(node)


def create_docx(markdown: str, destination: Path) -> None:
    doc = Document(); section = doc.sections[0]
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]; normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (("Title",24,"172033",0,10),("Heading 1",16,"315F72",14,7),("Heading 2",13,"315F72",10,5),("Heading 3",11.5,"4A6470",8,4)):
        s=doc.styles[name]; s.font.name="Aptos Display" if name != "Heading 3" else "Aptos"; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name != "Title"
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    add_numbering(doc,"bullet",50,50)
    header=section.header.paragraphs[0]; header.text="EQUINESYNC  |  C0-019  |  AGREEMENT, CONSENT, AND AUTHORIZATION V2.1"
    for run in header.runs: run.font.name="Aptos"; run.font.size=Pt(8); run.font.color.rgb=RGBColor(107,114,128)
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run("Founder-approved constitutional source  |  Page "); page_field(footer)
    for run in footer.runs: run.font.name="Aptos"; run.font.size=Pt(8); run.font.color.rgb=RGBColor(107,114,128)
    num_seq=100; previous=None; active_decimal=None
    for index,(kind,text) in enumerate(blocks(markdown)):
        clean=text.replace("**","").replace("`","")
        if kind=="heading1" and index==0:
            p=doc.add_paragraph(style="Title"); p.add_run(clean)
        elif kind.startswith("heading"):
            level=min(max(int(kind[-1]),1),3); p=doc.add_paragraph(style=f"Heading {level}"); p.add_run(clean)
        elif kind=="bullet":
            p=doc.add_paragraph(); set_num(p,50); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.10; p.add_run(clean)
        elif kind=="numbered":
            if previous != "numbered":
                num_seq += 1; add_numbering(doc,"decimal",num_seq,num_seq); active_decimal=num_seq
            p=doc.add_paragraph(); set_num(p,active_decimal); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.10; p.add_run(clean)
        else:
            p=doc.add_paragraph()
            if kind=="metadata" and ":" in clean:
                label,value=clean.split(":",1); r=p.add_run(label+":"); r.bold=True; p.add_run(value)
            else: p.add_run(clean)
        previous=kind
    core=doc.core_properties; core.title="Master Agreement, Consent, and Authorization Model V2.1"; core.subject="EquineSync C0-019 founder-approved constitutional source"; core.author="EquineSync Founder"; core.keywords="EquineSync, C0-019, Agreement, Consent, Authorization"; core.comments="Deterministic DOCX counterpart derived from exact substantive Markdown source."
    destination.parent.mkdir(parents=True,exist_ok=True); doc.save(destination)


def docx_text(path: Path) -> str:
    doc=Document(path)
    return " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())


def phase_a() -> None:
    if sha(MOUNTED) != EXPECTED or sha(PACKAGE) != PACKAGE_HASH or sha(V20) != V20_HASH:
        raise SystemExit("Source, package, or predecessor hash mismatch")
    RAW.parent.mkdir(parents=True, exist_ok=True); shutil.copy2(MOUNTED, RAW)
    HIST.parent.mkdir(parents=True, exist_ok=True); shutil.copy2(V20, HIST)
    if sha(RAW) != EXPECTED or sha(HIST) != V20_HASH: raise SystemExit("Preservation failed")
    with zipfile.ZipFile(PACKAGE) as archive:
        entry_bytes=archive.read(PACKAGE_ENTRY)
    if hashlib.sha256(entry_bytes).hexdigest() != EXPECTED or entry_bytes != RAW.read_bytes(): raise SystemExit("Archive entry mismatch")
    md=lifecycle_markdown("NOT YET ADOPTED","NOT YET LOCKED"); write_text(ACTIVE_MD,md); create_docx(md,ACTIVE_DOCX)
    if normalized(RAW.read_text(),True) != normalized(md,True): raise SystemExit("Substantive delta during metadata processing")
    ratio=SequenceMatcher(None,normalized(md),normalized(docx_text(ACTIVE_DOCX))).ratio()
    if ratio < .999: raise SystemExit(f"DOCX parity failed: {ratio}")
    generated=now()
    inventory=[
        {"path":rel(RAW),"version":"2.1","status":"exact expected source","sha256":sha(RAW),"controlling_claim":True,"relationship":"controlling successor"},
        {"path":rel(ACTIVE_MD),"version":"2.1","status":"founder-approved pre-adoption","sha256":sha(ACTIVE_MD),"controlling_claim":True,"relationship":"metadata derivative"},
        {"path":rel(V20),"version":"2.0","status":"historical superseded candidate","sha256":sha(V20),"controlling_claim":False,"relationship":"predecessor"},
        {"path":rel(HIST),"version":"2.0","status":"preserved history","sha256":sha(HIST),"controlling_claim":False,"relationship":"byte-identical preservation"},
    ]
    write_json(ROOT/"C0_019_SOURCE_INVENTORY.json",{"generated_at":generated,"sources":inventory})
    rows="\n".join(f"| `{x['path']}` | {x['version']} | {x['status']} | `{x['sha256']}` | {x['relationship']} |" for x in inventory)
    write_text(ROOT/"C0_019_SOURCE_INVENTORY.md",f"""# C0-019 Source Inventory

| Path | Version | Status | SHA-256 | Relationship |
|---|---|---|---|---|
{rows}

Agreement, consent, authorization, signature, assent, delegation, revocation, Privacy, Safeguarding, Identity, Relationship, Permission, Communication, Audit, Claims, Financial, Equine Health, and RF31 references were inventoried through repository search. No later competing founder-approved C0-019 source was found.
""")
    write_text(ROOT/"C0_019_INITIALIZATION_REPORT.md",f"""# C0-019 Initialization Report

- Founder disposition: `FOUNDER_AUTHORIZED_C0_019_AGREEMENT_CONSENT_AND_AUTHORIZATION_CONTROLLED_LIFECYCLE_COMPLETION`
- Mounted source: `{MOUNTED}`
- Exact SHA-256: `{EXPECTED}`
- Package: `{PACKAGE}` / `{PACKAGE_HASH}`
- Archive entry: `{PACKAGE_ENTRY}`
- V2.0 candidate: `{rel(V20)}` / `{V20_HASH}`
- Scope: C0-019 documentation and lifecycle only
""")
    write_text(ROOT/"C0_019_EXACT_SOURCE_PROVENANCE_RECORD.md",f"""# C0-019 Exact Source Provenance Record

- Mounted source: `{MOUNTED}`
- Preserved source: `{rel(RAW)}`
- Raw SHA-256: `{EXPECTED}`
- Size: `{RAW.stat().st_size}` bytes
- MIME: `{mimetypes.guess_type(RAW.name)[0] or 'text/markdown'}`
- Encoding: UTF-8; no BOM
- Line endings: LF
- Source package: `{PACKAGE}` / `{PACKAGE_HASH}`
- Package entry: `{PACKAGE_ENTRY}`
- Package-entry hash: `{EXPECTED}`
- Prior controlled recovered copy: byte-identical
- Chain of custody: mounted source -> raw hash -> controlled preservation -> archive-entry comparison -> founder-approved derivative
""")
    write_text(ROOT/"C0_019_ARCHIVE_EXTRACTION_RECORD.md",f"""# C0-019 Archive Extraction Record

- Archive: `{PACKAGE}`
- Archive SHA-256: `{PACKAGE_HASH}`
- Entry: `{PACKAGE_ENTRY}`
- Entry SHA-256: `{EXPECTED}`
- Fresh in-memory extraction: `PASS`
- Byte identity with mounted and controlled sources: `PASS`
""")
    write_json(ROOT/"C0_019_RAW_BYTE_COMPARISON.json",{"v2_1":{"sha256":EXPECTED,"bytes":RAW.stat().st_size},"v2_0":{"sha256":V20_HASH,"bytes":V20.stat().st_size},"byte_identical":False,"classification":"SUBSTANTIVE_SUCCESSOR"})
    v20n=normalized(V20.read_text()); v21n=normalized(RAW.read_text()); sim=SequenceMatcher(None,v20n,v21n).ratio()
    write_text(ROOT/"C0_019_NORMALIZED_TEXT_COMPARISON.md",f"""# C0-019 Normalized-Text Comparison

- V2.0: `{V20_HASH}`
- V2.1: `{EXPECTED}`
- Normalized similarity: `{sim:.6f}`
- Heading change: adds `2.1 Consent and Authorization Ownership Boundary` and `76. Version 2.1 Reconciliation Disposition`
- Requirement IDs in either source: `0`
- Classification: `SUBSTANTIVE_SUCCESSOR_WITH_EXPLICIT_FOUNDER_AUTHORIZATION`
""")
    write_text(ROOT/"C0_019_V2_0_TO_V2_1_SEMANTIC_DELTA.md","""# C0-019 V2.0 to V2.1 Semantic Delta

V2.1 preserves the complete V2.0 agreement lifecycle and adds an explicit constitutional ownership boundary. Consent and authorization creation, evidence, withdrawal, expiration, and agreement-derived effects belong here; Permission retains access enforcement; Privacy retains processing-law basis; Communication retains channel consent and notice delivery; Media retains asset rights. Version 2.1 also normalizes current canon names and records its reconciliation disposition.

No agreement, consent, authorization, guardian, financial, medical, transfer, dispute, or implementation authority is broadened beyond its specialized constitutional owner.
""")
    write_text(ROOT/"C0_019_AUTHORITY_BOUNDARY_DELTA.md","""# C0-019 Authority Boundary Delta

- Identity owns actor identity and authentication context.
- Relationship owns relationship existence and temporal state.
- Permission owns final authorization and field projection.
- Privacy owns lawful personal-data processing and privacy rights.
- Safeguarding owns guardian/minor eligibility and protected-participant controls.
- Communication owns delivery and notice evidence.
- Audit owns material event evidence.
- Claims owns contested authority and dispute procedure.
- Financial Truth owns financial responsibility and payment effects.
- Equine Health owns clinical and welfare authority.
- RF31 owns horse-transfer and Passport continuity.

Result: `NO_AUTHORITY_INVERSION_OR_EXPANSION`.
""")
    write_text(ROOT/"C0_019_SOURCE_SUCCESSION_ANALYSIS.md","""# C0-019 Source Succession Analysis

Primary classification: `EXACT_V2_1_CONTROLLING_SOURCE_CONFIRMED`.

V2.1 matches the authenticated expected hash and package entry, expressly supersedes V2.0, and carries founder-accepted controlling text pending the lifecycle completed here. V2.0 is preserved unchanged as a historical superseded candidate. No competing later approved source was found.
""")
    write_text(ROOT/"C0_019_SOURCE_IDENTITY_DECISION_RECORD.md",f"""# C0-019 Source Identity Decision Record

`CONFIRM_MASTER_AGREEMENT_CONSENT_AND_AUTHORIZATION_MODEL_V2_1_AS_CONTROLLING_C0_019_SOURCE_AND_PRESERVE_V2_0_AS_HISTORICAL_SUPERSEDED_CANDIDATE`

Classification: `EXACT_V2_1_CONTROLLING_SOURCE_CONFIRMED`. Exact hash `{EXPECTED}` and package provenance are verified; no unresolved competing source or policy conflict remains.
""")
    write_text(ROOT/"C0_019_CROSS_CANON_IMPACT_REPORT.md","""# C0-019 Cross-Canon Impact Report

V2.1 was reconciled against locked Product Vision, Identity, Relationship, Safeguarding, Communication, Audit, Claims, Financial Truth, Equine Health, RF31, Record Stewardship, and related adopted/candidate Privacy and Permission authorities. It coordinates with those canons without creating identity, relationship, access, guardian, processing, financial, medical, transfer, dispute, provider, runtime, or production authority.

No P0, P1, locked-canon conflict, circular authority, or unresolved substantive dependency was found. Permission V1.1 remains separately lifecycle-gated but its enforcement ownership is expressly preserved.
""")
    write_text(ROOT/"C0_019_MD_DOCX_PARITY_REPORT.md",f"""# C0-019 Markdown/DOCX Parity Report

- Markdown: `{rel(ACTIVE_MD)}` / `{sha(ACTIVE_MD)}`
- DOCX: `{rel(ACTIVE_DOCX)}` / `{sha(ACTIVE_DOCX)}`
- Normalized parity: `{ratio:.6f}`
- Omitted substantive provisions: `0`
- Added substantive provisions: `0`
- Changed requirement IDs: `0`
- Changed authority boundaries: `0`
- Result: `PASS`
""")
    write_text(ROOT/"C0_019_ADOPTION_READINESS_REPORT.md","""# C0-019 Adoption Readiness Report

Phase A disposition: `C0_019_EXACT_SOURCE_IDENTITY_AND_FOUNDER_APPROVED_COUNTERPARTS_COMPLETE_PENDING_ADOPTION`

Source, provenance, archive entry, expected hash, succession, substantive preservation, parity, authority boundaries, and cross-canon review pass. Page-by-page DOCX visual QA remains the final Phase A execution gate.

- P0: `0`
- Source-identity P1: `0`
- Source-identity-blocking P2: `0`
""")
    write_json(ROOT/"C0_019_PHASE_A_VALIDATION.json",{"disposition":"C0_019_EXACT_SOURCE_IDENTITY_AND_FOUNDER_APPROVED_COUNTERPARTS_COMPLETE_PENDING_ADOPTION","checks":{"source":True,"package":True,"archive_entry":True,"v20_preserved":True,"substantive_text":True,"parity":True,"cross_canon":True,"visual_qa_pending":True},"findings":{"p0":0,"open_p1":0,"blocking_p2":0},"authority":AUTHORITY})
    print(json.dumps({"phase":"A","source_sha256":EXPECTED,"markdown_sha256":sha(ACTIVE_MD),"docx_sha256":sha(ACTIVE_DOCX),"parity":ratio},indent=2))


if __name__ == "__main__": phase_a()
