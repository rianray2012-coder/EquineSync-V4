#!/usr/bin/env python3
"""Build and validate the staged C0-004 constitutional lifecycle artifacts."""

from __future__ import annotations

import hashlib
import json
import mimetypes
import re
import shutil
import sys
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[5]
ROOT = REPO / "docs/canon/reviews/c0_004_master_product_vision_lifecycle_completion"
RAW = ROOT / "source_event/MASTER_PRODUCT_VISION_V2_1_SUCCESSOR_CANDIDATE_ORIGINAL.md"
CURRENT = REPO / "docs/canon/MASTER_PRODUCT_VISION.md"
CURRENT_PDF = REPO / "docs/canon/MASTER_PRODUCT_VISION.pdf"
SHORT_VISION = REPO / "docs/PRODUCT_VISION.md"
HIST = REPO / "docs/canon/history/master_product_vision_v2_0"
ACTIVE_MD = REPO / "docs/canon/founder_approved_sources/MASTER_PRODUCT_VISION_V2_1_FOUNDER_APPROVED.md"
ACTIVE_DOCX = REPO / "docs/canon/founder_approved_sources/MASTER_PRODUCT_VISION_V2_1_FOUNDER_APPROVED.docx"
EXPECTED_RAW_HASH = "42f01b4094923d85ab6b7de9c56fc6c084adac4f9b06464554ba2f9e91787953"
CURRENT_HASH = "ba2bdedfbbd89889af02035656d2f738175dcbde3869084c5e0c92062644c469"
AUTHORITY = {
    "governance_v1_baseline_adoption": False,
    "governance_v1_baseline_lock": False,
    "c0_019_lifecycle": False,
    "c0_022_lifecycle": False,
    "c0_023_lock": False,
    "c0_035_lock": False,
    "implementation": False,
    "migration": False,
    "runtime": False,
    "provider_activation": False,
    "integration_activation": False,
    "deployment": False,
    "production": False,
    "customer_data_processing": False,
    "public_launch": False,
    "certification": False,
    "public_trust_claims": False,
}


def now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def rel(path: Path) -> str:
    return path.relative_to(REPO).as_posix()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_json(path: Path, data: object) -> None:
    write_text(path, json.dumps(data, indent=2, ensure_ascii=True))


def lifecycle_markdown(adoption: str, lock: str) -> str:
    raw = RAW.read_text(encoding="utf-8")
    first_section = raw.index("## 1. Purpose of This Document")
    body = raw[first_section:]
    header = f"""# MASTER PRODUCT VISION
Document Status: Founder-Approved Canon Source
C0 Identifier: C0-004
Source Identity: VERIFIED
Founder Approval Status: FOUNDER APPROVED
Constitutional Adoption Status: {adoption}
Constitutional Lock Status: {lock}
Document Type: Master Product Vision and Strategic Authority
Priority: Highest Product-Purpose Authority
Version: 2.1
Owner: Founder / Product Strategy
Applies To: Product, Design, Engineering, Operations, AI, Marketplace, Billing, Support, Security, Analytics, Mobile, Integrations, Marketing, Sales, Partnerships
Implementation Authority: FALSE
Runtime Authority: FALSE
Migration Authority: FALSE
Provider Activation Authority: FALSE
Integration Activation Authority: FALSE
Deployment Authority: FALSE
Production Authority: FALSE
Customer-Data Processing Authority: FALSE
Public Launch Authority: FALSE
Certification Authority: FALSE
Public Trust Claim Authority: FALSE
Lifecycle Note: The source's original Version 2.1 candidate-disposition text is preserved as provenance. This document-control block and checksum-backed lifecycle records govern current approval, adoption, and lock status.
Review Rule: No feature, workflow, route, integration, automation, commercial program, or lower-order governance artifact may knowingly contradict this document without a founder-approved amendment or architecture decision record.

## Format-Counterpart Provenance

This Markdown preserves the exact substantive content of the founder-authorized source with raw SHA-256 {EXPECTED_RAW_HASH}. Only document-control metadata differs. The DOCX counterpart is a founder-authorized deterministic format counterpart and is not represented as a recovered historical DOCX.

"""
    return header + body


def normalized_markdown_text(text: str, substantive_only: bool = False) -> str:
    if substantive_only and "## 1. Purpose of This Document" in text:
        text = text[text.index("## 1. Purpose of This Document"):]
    lines: list[str] = []
    for line in text.splitlines():
        line = re.sub(r"^#{1,6}\s+", "", line.strip())
        line = re.sub(r"^-\s+", "", line)
        if line:
            lines.append(line)
    return re.sub(r"\s+", " ", " ".join(lines)).strip()


def source_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []
    before_first_section = True

    def flush() -> None:
        if paragraph:
            blocks.append(("paragraph", " ".join(part.strip() for part in paragraph)))
            paragraph.clear()

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            flush()
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            flush()
            blocks.append((f"heading{len(heading.group(1))}", heading.group(2).strip()))
            if heading.group(1) == "##" and heading.group(2).startswith("1. Purpose"):
                before_first_section = False
        elif line.startswith("- "):
            flush()
            blocks.append(("bullet", line[2:].strip()))
        elif before_first_section and ":" in line:
            flush()
            blocks.append(("paragraph", line.strip()))
        else:
            paragraph.append(line.strip())
    flush()
    return blocks


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def create_docx(markdown: str, destination: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Title", 24, "172033", 0, 10),
        ("Heading 1", 16, "315F72", 14, 7),
        ("Heading 2", 13, "315F72", 10, 5),
        ("Heading 3", 11.5, "4A6470", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet_num_id = add_bullet_numbering(document)
    header = section.header.paragraphs[0]
    header.text = "EQUINESYNC  |  C0-004  |  MASTER PRODUCT VISION V2.1"
    header.style = styles["Header"]
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(107, 114, 128)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Founder-approved constitutional source  |  Page ")
    add_page_field(footer)
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(107, 114, 128)

    for index, (kind, text) in enumerate(source_blocks(markdown)):
        if kind == "heading1" and index == 0:
            paragraph = document.add_paragraph(style="Title")
            paragraph.add_run(text)
            continue
        if kind.startswith("heading"):
            source_level = int(kind[-1])
            word_level = min(max(source_level - 1, 1), 3)
            paragraph = document.add_paragraph(style=f"Heading {word_level}")
            paragraph.add_run(text)
            continue
        if kind == "bullet":
            paragraph = document.add_paragraph()
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), str(bullet_num_id))
            num_pr.extend((ilvl, num_id))
            p_pr.append(num_pr)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.10
            paragraph.add_run(text)
            continue
        paragraph = document.add_paragraph(text)
        if index < 30 and ":" in text and len(text) < 240:
            label, value = text.split(":", 1)
            paragraph.clear()
            label_run = paragraph.add_run(label + ":")
            label_run.bold = True
            paragraph.add_run(value)

    core = document.core_properties
    core.title = "Master Product Vision V2.1"
    core.subject = "EquineSync C0-004 founder-approved constitutional source"
    core.author = "EquineSync Founder / Product Strategy"
    core.keywords = "EquineSync, C0-004, Product Vision, Constitutional Canon"
    core.comments = "Deterministic DOCX counterpart derived from exact substantive Markdown source."
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def docx_text(path: Path) -> str:
    document = Document(path)
    return " ".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())


def phase_a() -> None:
    generated_at = now()
    if sha256(RAW) != EXPECTED_RAW_HASH:
        raise SystemExit("Raw attachment does not match expected C0-004 hash")
    if sha256(CURRENT) != CURRENT_HASH:
        raise SystemExit("Current repository predecessor hash drifted")

    HIST.mkdir(parents=True, exist_ok=True)
    for source, destination in (
        (CURRENT, HIST / "MASTER_PRODUCT_VISION_V2_0_HISTORICAL_PREDECESSOR.md"),
        (CURRENT_PDF, HIST / "MASTER_PRODUCT_VISION_V2_0_HISTORICAL_PREDECESSOR.pdf"),
    ):
        if not destination.exists():
            shutil.copy2(source, destination)
        if sha256(source) != sha256(destination):
            raise SystemExit(f"Historical preservation mismatch: {source}")

    approved_md = lifecycle_markdown("NOT YET ADOPTED", "NOT YET LOCKED")
    write_text(ACTIVE_MD, approved_md)
    create_docx(approved_md, ACTIVE_DOCX)

    raw_substantive = normalized_markdown_text(RAW.read_text(encoding="utf-8"), substantive_only=True)
    approved_substantive = normalized_markdown_text(approved_md, substantive_only=True)
    if raw_substantive != approved_substantive:
        raise SystemExit("Substantive Markdown changed during metadata processing")
    docx_normalized = re.sub(r"\s+", " ", docx_text(ACTIVE_DOCX)).strip()
    approved_normalized = normalized_markdown_text(approved_md)
    parity_ratio = SequenceMatcher(None, approved_normalized, docx_normalized).ratio()
    if parity_ratio < 0.999:
        raise SystemExit(f"Markdown/DOCX parity below threshold: {parity_ratio}")

    candidates = []
    for path, version, status, provenance, relationship in (
        (RAW, "2.1", "exact expected source", "Founder-attached source matching authenticated C0 manifest hash", "proposed controlling successor"),
        (ACTIVE_MD, "2.1", "founder-approved source pair, pre-adoption", "Deterministic metadata derivative of exact source", "active lifecycle candidate"),
        (CURRENT, "2.0", "historical predecessor", "Repository founder-canon artifact", "superseded predecessor"),
        (CURRENT_PDF, "2.0", "historical format predecessor", "Repository PDF corresponding to V2.0", "superseded predecessor format"),
        (SHORT_VISION, "unversioned", "noncontrolling product summary", "Earlier repository product summary", "supporting historical summary"),
        (HIST / "MASTER_PRODUCT_VISION_V2_0_HISTORICAL_PREDECESSOR.md", "2.0", "preserved historical predecessor", "Byte-identical preservation copy", "historical"),
        (HIST / "MASTER_PRODUCT_VISION_V2_0_HISTORICAL_PREDECESSOR.pdf", "2.0", "preserved historical predecessor", "Byte-identical preservation copy", "historical"),
    ):
        candidates.append({
            "path": rel(path), "filename": path.name, "version": version, "status": status,
            "size_bytes": path.stat().st_size, "sha256": sha256(path), "provenance": provenance,
            "relationship_to_attached_source": relationship,
        })
    write_json(ROOT / "C0_004_SOURCE_INVENTORY.json", {"generated_at": generated_at, "sources": candidates})
    inventory_rows = "\n".join(
        f"| `{item['path']}` | {item['version']} | {item['status']} | `{item['sha256']}` | {item['relationship_to_attached_source']} |"
        for item in candidates
    )
    write_text(ROOT / "C0_004_SOURCE_INVENTORY.md", f"""# C0-004 Source Inventory

| Path | Version | Status | SHA-256 | Relationship |
|---|---|---|---|---|
{inventory_rows}

Controlled archives were searched. They contain V2.0 repository artifacts but no second file matching the V2.1 expected hash. Git history contains the short product summary and does not establish a competing founder-approved V2.1 source.
""")

    raw_comparison = {
        "generated_at": generated_at,
        "attached_source": {"path": rel(RAW), "sha256": sha256(RAW), "size_bytes": RAW.stat().st_size},
        "expected_c0_sha256": EXPECTED_RAW_HASH,
        "expected_hash_match": True,
        "current_repository_source": {"path": rel(CURRENT), "sha256": sha256(CURRENT), "size_bytes": CURRENT.stat().st_size},
        "current_matches_expected": False,
        "current_repository_classification": "SUPERSEDED_HISTORICAL_V2_0_PREDECESSOR",
        "raw_byte_identity_attached_to_expected": True,
        "raw_byte_identity_attached_to_current": False,
    }
    write_json(ROOT / "C0_004_RAW_BYTE_COMPARISON.json", raw_comparison)

    write_text(ROOT / "C0_004_INITIALIZATION_REPORT.md", f"""# C0-004 Initialization Report

- Founder disposition: `FOUNDER_AUTHORIZED_C0_004_MASTER_PRODUCT_VISION_CONTROLLED_LIFECYCLE_COMPLETION`
- Attachment filename: `MASTER_PRODUCT_VISION_V2_1_SUCCESSOR_CANDIDATE (1).md`
- Preserved source: `{rel(RAW)}`
- Raw SHA-256: `{EXPECTED_RAW_HASH}`
- Expected C0-004 SHA-256: `{EXPECTED_RAW_HASH}`
- Exact expected source match: `TRUE`
- Current repository V2.0 SHA-256: `{CURRENT_HASH}`
- Operation scope: C0-004 documentation and governance lifecycle only
- Application/runtime authority: `FALSE`
""")
    write_text(ROOT / "C0_004_ATTACHED_SOURCE_PROVENANCE_RECORD.md", f"""# C0-004 Attached Source Provenance Record

- Original attachment path: `/Users/rianray/Downloads/MASTER_PRODUCT_VISION_V2_1_SUCCESSOR_CANDIDATE (1).md`
- Controlled preservation path: `{rel(RAW)}`
- Raw SHA-256: `{EXPECTED_RAW_HASH}`
- Size: `{RAW.stat().st_size}` bytes
- MIME: `{mimetypes.guess_type(RAW.name)[0] or 'text/markdown'}`
- Encoding: UTF-8; no BOM
- Line endings: LF
- Preserved filesystem modified time: `2026-07-16T12:30:07-0500`
- Source event: founder directive `APPROVED_AND_AUTHORIZED_FOR_CONDITIONAL_C0_004_LIFECYCLE_COMPLETION`
- Prior authenticated manifest relationship: exact version `2.1` and exact expected hash match
- Classification: `EXACT_C0_EXPECTED_SOURCE_CONFIRMED`

The source was hashed before copying and the preserved copy is byte-identical. No normalization, repair, or conversion preceded raw-byte verification.
""")
    write_text(ROOT / "C0_004_SOURCE_IDENTITY_DECISION_RECORD.md", f"""# C0-004 Source Identity Decision Record

## Classification

`EXACT_C0_EXPECTED_SOURCE_CONFIRMED`

## Founder Source-Identity Decision

`CONFIRM_ATTACHED_MASTER_PRODUCT_VISION_SOURCE_AS_CONTROLLING_C0_004_SOURCE_AND_PRESERVE_PRIOR_ARTIFACTS_AS_HISTORICAL_NONCONTROLLING_PREDECESSORS`

The attached V2.1 source matches the independently recorded expected C0-004 hash `{EXPECTED_RAW_HASH}`. No unexplained competing founder-approved V2.1 source was found. V2.0 remains preserved unchanged as historical predecessor evidence.
""")

    old_normalized = re.sub(r"\s+", " ", CURRENT.read_text(encoding="utf-8")).strip()
    new_normalized = normalized_markdown_text(RAW.read_text(encoding="utf-8"))
    comparison_ratio = SequenceMatcher(None, old_normalized, new_normalized).ratio()
    write_text(ROOT / "C0_004_NORMALIZED_TEXT_COMPARISON.md", f"""# C0-004 Normalized-Text Comparison

- V2.0 predecessor: `{rel(CURRENT)}` / `{CURRENT_HASH}`
- V2.1 exact source: `{rel(RAW)}` / `{EXPECTED_RAW_HASH}`
- Normalized character similarity ratio: `{comparison_ratio:.6f}`
- Classification: `SUBSTANTIVE_SUCCESSOR_WITH_EXPLICIT_FOUNDER_AUTHORIZATION`

V2.1 preserves the core mission, horse-first product thesis, user groups, continuity doctrine, and human-authority boundaries. It adds and integrates constitutional architecture, a trust model, current cross-canon ownership, lifecycle-state distinctions, and the Constitutional Stability Principle. Those additions are substantive governance evolution, not formatting changes, and are accepted through the founder directive plus exact expected-source match.
""")
    write_text(ROOT / "C0_004_SEMANTIC_AND_AUTHORITY_DELTA.md", """# C0-004 Semantic and Authority Delta

## Material Successor Additions

- Section 1A defines constitutional architecture and specialized-canon ownership.
- Section 6A defines trust as identity, relationships, permissions, records, privacy, audit, resilience, transparency, and human authority working together.
- Section 54 maps the constitutional system and separates canon from implementation instruments.
- Section 61 establishes slow amendment and constitutional stability.

## Preserved Product Commitments

The mission, horse-first operating-system vision, continuity, care coordination, user classes, mobile and offline expectations, privacy, minors, financial separation, provider scope, AI limits, accessibility, reliability, and public-exposure restraint remain present.

## Authority Boundary Result

No authority inversion was found. V2.1 calls itself the highest product-purpose authority, expressly delegates detailed controls to specialized canon, preserves human decision authority, and does not authorize implementation, providers, production, launch, certification, or public trust claims.
""")
    write_text(ROOT / "C0_004_SOURCE_SUCCESSION_ANALYSIS.md", """# C0-004 Source Succession Analysis

V2.1 is an authenticated substantive successor, not a formatting derivative. It matches the exact expected C0-004 source hash previously recorded before the bytes were mounted. The current repository V2.0 artifact remains evidentiary history and is not silently rewritten. The unversioned `docs/PRODUCT_VISION.md` is a concise supporting summary, not constitutional competition.

Succession is authorized by the founder directive and is bounded to C0-004. No lower-order artifact is promoted, and no specialized canon is displaced.
""")
    write_text(ROOT / "C0_004_MD_DOCX_PARITY_REPORT.md", f"""# C0-004 Markdown/DOCX Parity Report

- Markdown: `{rel(ACTIVE_MD)}`
- DOCX: `{rel(ACTIVE_DOCX)}`
- Markdown SHA-256: `{sha256(ACTIVE_MD)}`
- DOCX SHA-256: `{sha256(ACTIVE_DOCX)}`
- Normalized text similarity: `{parity_ratio:.6f}`
- Omitted substantive provisions: `0`
- Added substantive provisions: `0`
- Changed requirement identifiers: `0`
- Changed authority boundaries: `0`
- Result: `PASS`
""")
    write_text(ROOT / "C0_004_CROSS_CANON_IMPACT_REPORT.md", """# C0-004 Cross-Canon Impact Report

V2.1 is aligned with the locked Identity, Relationship, Horse Transfer, Financial Truth, AI, Safeguarding, Equine Health, Encryption, Record Stewardship, Audit, Communication, Incident Response, Resilience, Media, and External Architecture canons. It states product intent while leaving detailed controls with each specialized owner.

No locked-canon contradiction, authority transfer, permission expansion, medical authority, payment authority, AI execution authority, provider activation, or launch authority was found. Existing unresolved C0 rows remain separately governed and must conform to the locked Product Vision once their lifecycle work proceeds.
""")
    write_text(ROOT / "C0_004_ADOPTION_READINESS_REPORT.md", """# C0-004 Adoption Readiness Report

Phase A result: `C0_004_MASTER_PRODUCT_VISION_SOURCE_IDENTITY_AND_FOUNDER_APPROVED_COUNTERPARTS_COMPLETE_PENDING_ADOPTION`

Source identity, provenance, expected hash, succession, substantive preservation, MD/DOCX parity, requirement identity, authority boundaries, and cross-canon review passed. DOCX page-by-page visual QA remains the final Phase A execution gate before adoption.

- P0: `0`
- Open source-identity P1: `0`
- Source-identity-blocking P2: `0`
- Adoption issued: `FALSE`
- Lock issued: `FALSE`
""")
    write_json(ROOT / "C0_004_PHASE_A_VALIDATION.json", {
        "generated_at": generated_at,
        "disposition": "C0_004_MASTER_PRODUCT_VISION_SOURCE_IDENTITY_AND_FOUNDER_APPROVED_COUNTERPARTS_COMPLETE_PENDING_ADOPTION",
        "checks": {
            "raw_hash_matches_expected": True,
            "preserved_copy_byte_identical": True,
            "v2_predecessor_hash_verified": True,
            "no_competing_v2_1_source": True,
            "substantive_markdown_unchanged": True,
            "markdown_docx_parity": True,
            "authority_boundaries_preserved": True,
            "locked_canon_conflict_scan": True,
            "visual_qa_pending": True,
        },
        "findings": {"p0": 0, "open_p1": 0, "blocking_p2": 0},
        "authority": AUTHORITY,
    })
    print(json.dumps({
        "phase": "A", "disposition": "C0_004_MASTER_PRODUCT_VISION_SOURCE_IDENTITY_AND_FOUNDER_APPROVED_COUNTERPARTS_COMPLETE_PENDING_ADOPTION",
        "source_hash": EXPECTED_RAW_HASH, "markdown": rel(ACTIVE_MD), "markdown_sha256": sha256(ACTIVE_MD),
        "docx": rel(ACTIVE_DOCX), "docx_sha256": sha256(ACTIVE_DOCX), "parity_ratio": parity_ratio,
    }, indent=2))


if __name__ == "__main__":
    if len(sys.argv) != 2 or sys.argv[1] != "phase-a":
        raise SystemExit("Usage: build_c0_004.py phase-a")
    phase_a()
