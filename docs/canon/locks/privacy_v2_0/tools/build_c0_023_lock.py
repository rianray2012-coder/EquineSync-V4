#!/usr/bin/env python3
"""Lock C0-023 after its dependency rerun; retain the C0-035 blocker."""

from __future__ import annotations

import hashlib
import json
import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


REPO = Path(__file__).resolve().parents[5]
ROOT = REPO / "docs/canon/locks/privacy_v2_0"
READINESS = REPO / "docs/canon/adoptions/c0_023_c0_035_lock_readiness_rerun"
PRIVACY_MD = REPO / "docs/canon/founder_approved_sources/MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.md"
PRIVACY_DOCX = REPO / "docs/canon/founder_approved_sources/MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.docx"
ADOPTION = REPO / "docs/canon/adoptions/c0_023_privacy_v2_0/C0_023_PRIVACY_AND_DATA_PROTECTION_CONSTITUTIONAL_ADOPTION_RECORD.md"
ADOPTION_HASH = "1ba284afda6c02ef5dc1cf4ae3b7b9b963c5aeb223341fe95bc587b411ed9b12"
PRE_MD_HASH = "b773389d83308ab051e4959a4f61304cc3fb5328a5f8086ef683ab0a914498bb"
PRE_DOCX_HASH = "303b9ab9092a18deafc78d728c32f17bf8752e5383f5f793c0c47202bd305c7c"


def now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def rel(path: Path) -> str:
    return path.relative_to(REPO).as_posix()


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_json(path: Path, value: object) -> None:
    write_text(path, json.dumps(value, indent=2, ensure_ascii=True))


def replace_once(text: str, old: str, new: str) -> str:
    if text.count(old) != 1:
        raise RuntimeError(f"Expected exactly one metadata occurrence: {old}")
    return text.replace(old, new, 1)


def update_docx_paragraph(document: Document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one DOCX metadata paragraph: {old}")
    paragraph = matches[0]
    for run in paragraph.runs:
        run.text = ""
    paragraph.add_run(new)


def replace_table_row(path: Path, row_id: str, replacement: str) -> None:
    lines = path.read_text().splitlines()
    for index, line in enumerate(lines):
        if line.startswith(f"| {row_id} |"):
            lines[index] = replacement
            write_text(path, "\n".join(lines))
            return
    raise RuntimeError(f"Missing {row_id} in {path}")


def preflight() -> None:
    assert sha(PRIVACY_MD) == PRE_MD_HASH
    assert sha(PRIVACY_DOCX) == PRE_DOCX_HASH
    assert sha(ADOPTION) == ADOPTION_HASH
    state = json.loads((REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_CURRENT_LIFECYCLE_STATE_REGISTER.json").read_text())
    rows = {row["record_id"]: row for row in state["rows"]}
    assert rows["C0-004"]["lock_state"] == "LOCKED"
    assert rows["C0-019"]["lock_state"] == "LOCKED"
    assert rows["C0-022"]["current_exact_source_status"] == "EXACT_V1_1_CONTROLLING_SOURCE_FOUNDER_CONFIRMED"
    assert rows["C0-023"]["adoption_state"] == "ADOPTED" and rows["C0-023"]["lock_state"] == "NOT_YET_LOCKED"


def create_readiness_reports() -> None:
    READINESS.mkdir(parents=True, exist_ok=True)
    write_text(READINESS / "C0_023_PRIVACY_LOCK_READINESS_RERUN.md", """# C0-023 Privacy Lock-Readiness Rerun

Disposition: `C0_023_PRIVACY_LOCK_READINESS_PASSED`

- C0-004 Product Vision: `LOCKED`
- C0-019 Agreement, Consent, and Authorization: `LOCKED`
- Adopted Privacy source and adoption-record hashes: `VERIFIED`
- P0: `0`
- Open P1: `0`
- Lock-blocking P2: `0`
- Retained nonblocking P2: `C0-023-LR-P2-01`

`C0-023-LR-P1-01` is closed. The C0-028 and C0-041 change-impact observation remains assigned and nonblocking; their later adoption must conform to locked Privacy authority.
""")
    write_text(READINESS / "C0_035_REPORTING_LOCK_READINESS_RERUN.md", """# C0-035 Reporting Lock-Readiness Rerun

Disposition: `C0_035_REPORTING_LOCK_READINESS_COMPLETE_NOT_READY`

- C0-004 Product Vision: `LOCKED`
- C0-023 Privacy: `LOCKED_BY_THIS_SEQUENCED_DIRECTIVE`
- C0-022 Permission and Access-Control: `SOURCE_CONFIRMED; NOT_ADOPTED; NOT_LOCKED`
- P0: `0`
- Open P1: `1`
- Lock-blocking P2: `0`
- Retained nonblocking P2: `C0-035-LR-P2-01`

Open finding `C0-035-LR-P1-01` remains narrowed to C0-022 constitutional lifecycle completion. Founder lock authority does not waive the existing dependency gate. No C0-035 lock is issued.
""")
    write_json(READINESS / "C0_023_C0_035_LOCK_READINESS_RERUN_FINDINGS.json", {
        "C0-023": {"ready": True, "p0": 0, "open_p1": 0, "blocking_p2": 0, "retained_p2": ["C0-023-LR-P2-01"]},
        "C0-035": {"ready": False, "p0": 0, "open_p1": 1, "blocking_p2": 0, "retained_p2": ["C0-035-LR-P2-01"], "blocking_finding": "C0-035-LR-P1-01"},
    })


def lock_privacy() -> tuple[str, str]:
    pre_dir = ROOT / "pre_lock"
    pre_dir.mkdir(parents=True, exist_ok=True)
    for source in (PRIVACY_MD, PRIVACY_DOCX):
        destination = pre_dir / source.name
        shutil.copy2(source, destination)
        assert sha(source) == sha(destination)

    text = PRIVACY_MD.read_text()
    text = replace_once(text, "**FOUNDER APPROVED - CONSTITUTIONALLY ADOPTED - NOT YET LOCKED**", "**FOUNDER APPROVED - CONSTITUTIONALLY ADOPTED AND LOCKED**")
    text = replace_once(text, "Constitutional Status: FOUNDER-APPROVED ADOPTED CANON", "Constitutional Status: FOUNDER-APPROVED ADOPTED AND LOCKED CANON")
    text = replace_once(text, "Constitutional Lock Status: NOT YET LOCKED", "Constitutional Lock Status: LOCKED")
    write_text(PRIVACY_MD, text)

    document = Document(PRIVACY_DOCX)
    update_docx_paragraph(document, "FOUNDER APPROVED - CONSTITUTIONALLY ADOPTED - NOT YET LOCKED", "FOUNDER APPROVED - CONSTITUTIONALLY ADOPTED AND LOCKED")
    update_docx_paragraph(document, "Constitutional Status: FOUNDER-APPROVED ADOPTED CANON", "Constitutional Status: FOUNDER-APPROVED ADOPTED AND LOCKED CANON")
    update_docx_paragraph(document, "Constitutional Lock Status: NOT YET LOCKED", "Constitutional Lock Status: LOCKED")
    document.save(PRIVACY_DOCX)

    pre_md = (pre_dir / PRIVACY_MD.name).read_text()
    assert pre_md.split("## Authority Boundary", 1)[1] == PRIVACY_MD.read_text().split("## Authority Boundary", 1)[1]
    def body(path: Path) -> list[str]:
        texts = [paragraph.text for paragraph in Document(path).paragraphs]
        return texts[texts.index("Authority Boundary"):]
    assert body(pre_dir / PRIVACY_DOCX.name) == body(PRIVACY_DOCX)
    return sha(PRIVACY_MD), sha(PRIVACY_DOCX)


def create_lock_record(md_hash: str, docx_hash: str) -> tuple[Path, str]:
    record = ROOT / "C0_023_PRIVACY_AND_DATA_PROTECTION_CONSTITUTIONAL_LOCK_RECORD.md"
    write_text(record, f"""# C0-023 Privacy and Data Protection Constitutional Lock Record

Disposition: `C0_023_PRIVACY_AND_DATA_PROTECTION_CONSTITUTIONALLY_LOCKED`

- C0 identifier: `C0-023`
- Canon: `Master Privacy and Data Protection Model V2.0`
- Locked Markdown: `{rel(PRIVACY_MD)}` / `{md_hash}`
- Locked DOCX counterpart: `{rel(PRIVACY_DOCX)}` / `{docx_hash}`
- Adoption record: `{rel(ADOPTION)}` / `{ADOPTION_HASH}`
- Lock-readiness rerun: `{rel(READINESS / 'C0_023_PRIVACY_LOCK_READINESS_RERUN.md')}`
- Founder directive: `Lock 23 & 35`
- Lock timestamp: `{now()}`
- P0: `0`
- Open P1: `0`
- Lock-blocking P2: `0`
- Retained nonblocking P2: `C0-023-LR-P2-01`
- Constitutional lock: `LOCKED`
- Implementation, migration, runtime, provider, integration, deployment, production, customer-data-processing, public-launch, certification, and public-trust authority: `FALSE`

Later C0-028 and C0-041 lifecycle work must conform to this locked Privacy authority. Amendment or supersession requires a separate founder-authorized constitutional lifecycle event.
""")
    digest = sha(record)
    write_text(record.with_suffix(".sha256"), f"{digest}  {record.name}")
    return record, digest


def update_json_register(path: Path, collection: str, md_hash: str, record: Path) -> None:
    data = json.loads(path.read_text())
    row = next(item for item in data[collection] if item.get("record_id") == "C0-023")
    row.update({
        "current_category": "lifecycle_verified_complete",
        "current_exact_source_status": "EXACT_LOCKED_SOURCE_VERIFIED",
        "current_sha256": md_hash,
        "founder_status": "FOUNDER_APPROVED",
        "adoption_state": "ADOPTED",
        "lock_state": "LOCKED",
        "lock_readiness_state": "PASSED",
        "lifecycle_evidence": rel(record),
        "retained_p2": ["C0-023-LR-P2-01"],
        "unresolved_blocker": False,
        "remedy": "Complete; future changes require governed amendment or supersession.",
    })
    if "unresolved_lifecycle_blockers" in data:
        data["unresolved_lifecycle_blockers"] -= 1
        data["category_counts"]["lock_required"] -= 1
        data["category_counts"]["lifecycle_verified_complete"] += 1
    data["status"] = "UPDATED_AFTER_C0_023_CONSTITUTIONAL_LOCK_C0_035_REMAINS_BLOCKED"
    data["generated_at"] = now()
    write_json(path, data)


def update_governance(md_hash: str, record: Path, record_hash: str) -> None:
    batch = REPO / "docs/canon/reviews/governance_v1_0_c0_lifecycle_completion_batch_1"
    lifecycle = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle"
    update_json_register(batch / "C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.json", "rows", md_hash, record)
    update_json_register(lifecycle / "C0_CURRENT_LIFECYCLE_STATE_REGISTER.json", "rows", md_hash, record)

    source_md = batch / "C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.md"
    state_md = lifecycle / "C0_CURRENT_LIFECYCLE_STATE_REGISTER.md"
    replace_table_row(source_md, "C0-023", f"| C0-023 | Master Privacy and Data Protection Model | lifecycle_verified_complete | EXACT_LOCKED_SOURCE_VERIFIED | FOUNDER_APPROVED | ADOPTED | LOCKED | `{rel(PRIVACY_MD)}` |")
    replace_table_row(state_md, "C0-023", f"| C0-023 | Master Privacy and Data Protection Model | EXACT_LOCKED_SOURCE_VERIFIED | ADOPTED | LOCKED | `{rel(record)}` | False |")
    write_text(source_md, source_md.read_text().replace("Status: `UPDATED_AFTER_C0_022_V1_1_SOURCE_IDENTITY_CONFIRMATION`", "Status: `UPDATED_AFTER_C0_023_CONSTITUTIONAL_LOCK_C0_035_REMAINS_BLOCKED`"))
    write_text(state_md, state_md.read_text().replace("unresolved lifecycle blockers: `15`", "unresolved lifecycle blockers: `14`"))

    ledger_json = lifecycle / "C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.json"
    ledger = json.loads(ledger_json.read_text())
    ledger["records"] = [row for row in ledger["records"] if row["record_id"] != "C0-023"]
    ledger["row_count"] = len(ledger["records"])
    ledger["track_counts"]["TRACK_D_MISSING_SOURCE_RECOVERY"] -= 1
    ledger["status"] = "UPDATED_AFTER_C0_023_CONSTITUTIONAL_LOCK_C0_035_REMAINS_BLOCKED"
    ledger["purpose"] = "Resolution planning for the 14 remaining blockers; C0-023 completed separately."
    ledger["generated_at"] = now()
    write_json(ledger_json, ledger)
    for name in ("C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.md", "C0_UNRESOLVED_LIFECYCLE_BLOCKER_LEDGER.md"):
        path = lifecycle / name
        text = "\n".join(line for line in path.read_text().splitlines() if not line.startswith("| C0-023 |")) + "\n"
        text = text.replace("remaining lifecycle blockers: `15`", "remaining lifecycle blockers: `14`")
        text = text.replace("Status: `UPDATED_AFTER_C0_022_V1_1_SOURCE_IDENTITY_CONFIRMATION`", "Status: `UPDATED_AFTER_C0_023_CONSTITUTIONAL_LOCK_C0_035_REMAINS_BLOCKED`")
        write_text(path, text)

    delta_json = lifecycle / "C0_HISTORICAL_TO_CURRENT_DELTA.json"
    delta = json.loads(delta_json.read_text())
    row = next(item for item in delta["rows"] if item["record_id"] == "C0-023")
    row.update({"current_lock": "LOCKED", "current_category": "lifecycle_verified_complete", "current_founder_status": "FOUNDER_APPROVED", "basis": rel(record)})
    delta["generated_at"] = now()
    write_json(delta_json, delta)
    replace_table_row(lifecycle / "C0_HISTORICAL_TO_CURRENT_DELTA.md", "C0-023", f"| C0-023 | UNRESOLVED | FOUNDER_APPROVED | UNRESOLVED | LOCKED | `{rel(record)}` |")

    index = REPO / "docs/canon/CANON_INDEX.md"
    text = index.read_text()
    old = next(line for line in text.splitlines() if line.startswith("| 3 | Master Privacy and Data Protection Model |"))
    new = f"| 3 | Master Privacy and Data Protection Model | `{rel(PRIVACY_MD)}` | `ADOPTED`; `LOCKED`; retained nonblocking `C0-023-LR-P2-01` | Privacy and data-protection constitutional authority. No implementation, runtime, migration, provider, production, launch, certification, or public-trust authority. |"
    write_text(index, text.replace(old, new))

    lock_ledger = REPO / "docs/canon/registries/CANON_LOCK_LEDGER.md"
    text = lock_ledger.read_text()
    row_text = f"| Master Privacy and Data Protection Model | 2.0 | `{rel(PRIVACY_MD)}` | `{md_hash}` | `ADOPTED` | `LOCKED` | P0 0; P1 0; `C0-023-LR-P2-01` retained nonblocking | C0-023 lock record `{record_hash}` and readiness rerun | No implementation, migration, runtime, provider, production, launch, certification, or public-trust authority |"
    marker = "\nAny future substantive amendment requires"
    write_text(lock_ledger, text.replace(marker, "\n" + row_text + marker))

    for path in (REPO / "docs/canon/companions/GLOBAL_LOCK_STATE_REGISTER.json", REPO / "docs/canon/companions/GLOBAL_ADOPTION_STATE_REGISTER.json"):
        data = json.loads(path.read_text())
        item = next(entry for entry in data["entries"] if entry["artifact"] == "Master Privacy and Data Protection Model V2.0")
        item.update({"lock_state": "locked", "evidence": rel(record), "evidence_sha256": record_hash})
        data["generated_at"] = now()
        write_json(path, data)
    replace_table_row(REPO / "docs/canon/companions/GLOBAL_LOCK_STATE_REGISTER.md", "Master Privacy and Data Protection Model V2.0", f"| Master Privacy and Data Protection Model V2.0 | locked | {rel(record)} |")
    replace_table_row(REPO / "docs/canon/companions/GLOBAL_ADOPTION_STATE_REGISTER.md", "Master Privacy and Data Protection Model V2.0", f"| Master Privacy and Data Protection Model V2.0 | adopted; locked | {rel(record)} |")


def main() -> None:
    preflight()
    create_readiness_reports()
    md_hash, docx_hash = lock_privacy()
    record, record_hash = create_lock_record(md_hash, docx_hash)
    update_governance(md_hash, record, record_hash)
    write_json(ROOT / "C0_023_LOCK_BUILD_RECORD.json", {
        "disposition": "C0_023_PRIVACY_AND_DATA_PROTECTION_CONSTITUTIONALLY_LOCKED",
        "C0-023": {"locked": True, "markdown_sha256": md_hash, "docx_sha256": docx_hash, "lock_record_sha256": record_hash, "retained_p2": ["C0-023-LR-P2-01"]},
        "C0-035": {"locked": False, "open_p1": ["C0-035-LR-P1-01"], "blocker": "C0-022 is not adopted or locked"},
        "authority": {key: False for key in ("implementation", "migration", "runtime", "provider_activation", "production", "public_launch", "certification", "public_trust_claims")},
    })
    print(json.dumps({"C0-023": "LOCKED", "C0-035": "BLOCKED", "privacy_markdown_sha256": md_hash, "privacy_docx_sha256": docx_hash, "lock_record_sha256": record_hash}, indent=2))


if __name__ == "__main__":
    main()
