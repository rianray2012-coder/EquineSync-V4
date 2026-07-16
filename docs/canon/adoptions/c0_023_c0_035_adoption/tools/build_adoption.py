#!/usr/bin/env python3
"""Adopt C0-023 and C0-035 without granting lock or operational authority."""

from __future__ import annotations

import hashlib
import importlib.util
import json
import re
import shutil
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


REPO = Path(__file__).resolve().parents[5]
ROOT = REPO / "docs/canon/adoptions/c0_023_c0_035_adoption"
PRIVACY_DIR = REPO / "docs/canon/adoptions/c0_023_privacy_v2_0"
REPORTING_DIR = REPO / "docs/canon/adoptions/c0_035_reporting_v2_0"
APPROVED = REPO / "docs/canon/founder_approved_sources"
COMPLETION = REPO / "docs/canon/reviews/privacy_reporting_format_counterpart_completion_v1_0"

PRIVACY_MD = APPROVED / "MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.md"
PRIVACY_DOCX = APPROVED / "MASTER_PRIVACY_AND_DATA_PROTECTION_MODEL_V2_0_FOUNDER_APPROVED.docx"
REPORTING_MD = APPROVED / "MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0_FOUNDER_APPROVED.md"
REPORTING_DOCX = APPROVED / "MASTER_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_MODEL_V2_0_FOUNDER_APPROVED.docx"

PRE_HASHES = {
    PRIVACY_MD: "41d8a72e78c369a297350b0abbdda90b1bf5401a3bb16050101679c325012eed",
    PRIVACY_DOCX: "76cdc1338a70c1c18f0295340f0f3599c4e99b3d37c0f306d740db9537ef1846",
    REPORTING_MD: "56de16098bbba31571c94d2ec3560dcc48a74daae318da047e5b5592044e953a",
    REPORTING_DOCX: "d6cc3cc2add501d2d34d947de4e7a2ee7c6af77aa9669606c95514361657504d",
}
PACKAGE = REPO / "outputs/privacy_reporting_format_counterpart_completion_package.zip"
PACKAGE_SHA = "9316affecd7f82ed0604f164edfc07d311b8c61481c5c192c288c3582303c593"
ADOPTION_TS = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")

AUTHORITY = {
    "constitutional_lock": False,
    "governance_v1_baseline_adoption": False,
    "governance_v1_baseline_lock": False,
    "implementation": False,
    "migration": False,
    "runtime": False,
    "provider_activation": False,
    "integration_activation": False,
    "deployment": False,
    "production": False,
    "customer_data_processing": False,
    "public_launch": False,
    "certification": False,
    "public_trust_claims": False,
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def rel(path: Path) -> str:
    return path.relative_to(REPO).as_posix()


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_json(path: Path, data: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=2) + "\n", encoding="utf-8")


def replace_once(text: str, old: str, new: str) -> str:
    if text.count(old) != 1:
        raise ValueError(f"Expected one occurrence of {old!r}, found {text.count(old)}")
    return text.replace(old, new, 1)


def update_paragraph(document: Document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == old]
    if len(matches) != 1:
        raise ValueError(f"Expected one DOCX paragraph {old!r}, found {len(matches)}")
    paragraph = matches[0]
    for run in paragraph.runs:
        run.text = ""
    paragraph.add_run(new)


def insert_after(paragraph, text: str) -> None:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    created.style = paragraph.style
    created.add_run(text)


def preserve_pre_adoption() -> None:
    targets = {
        PRIVACY_MD: PRIVACY_DIR / "pre_adoption" / PRIVACY_MD.name,
        PRIVACY_DOCX: PRIVACY_DIR / "pre_adoption" / PRIVACY_DOCX.name,
        REPORTING_MD: REPORTING_DIR / "pre_adoption" / REPORTING_MD.name,
        REPORTING_DOCX: REPORTING_DIR / "pre_adoption" / REPORTING_DOCX.name,
    }
    for source, destination in targets.items():
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)
        if sha256(destination) != PRE_HASHES[source]:
            raise ValueError(f"Preservation hash mismatch: {destination}")


def preflight() -> None:
    if sha256(PACKAGE) != PACKAGE_SHA:
        raise SystemExit("Prerequisite evidence-package hash mismatch")
    for path, digest in PRE_HASHES.items():
        if sha256(path) != digest:
            raise SystemExit(f"Pre-adoption hash mismatch: {path}")
    manifest = json.loads((COMPLETION / "PRIVACY_REPORTING_FORMAT_COUNTERPART_MANIFEST.json").read_text(encoding="utf-8"))
    validation = json.loads((COMPLETION / "PRIVACY_REPORTING_FORMAT_COUNTERPART_VALIDATION.json").read_text(encoding="utf-8"))
    if not validation["checks"]["all_required_checks_pass"]:
        raise SystemExit("Prerequisite validation is not passing")
    if manifest["privacy"]["source_substantive_body_exact_match"] is not True:
        raise SystemExit("Privacy source parity missing")
    if manifest["reporting"]["approved_docx_byte_identical_to_source"] is not True:
        raise SystemExit("Reporting source identity missing")
    if any(manifest["authority"].values()):
        raise SystemExit("Prerequisite package contains authority overclaim")
    for path in PRE_HASHES:
        if path.suffix == ".md" and re.search(r"^(?:<{7}|={7}|>{7})", path.read_text(encoding="utf-8"), re.M):
            raise SystemExit(f"Merge marker found: {path}")


def update_privacy() -> None:
    text = PRIVACY_MD.read_text(encoding="utf-8")
    text = replace_once(text, "**FOUNDER APPROVED - FOUNDER-APPROVED CANON SOURCE - NOT YET ADOPTED - NOT YET LOCKED**", "**FOUNDER APPROVED - CONSTITUTIONALLY ADOPTED - NOT YET LOCKED**")
    text = replace_once(text, "Constitutional Status: FOUNDER-APPROVED CANON SOURCE", "Constitutional Status: FOUNDER-APPROVED ADOPTED CANON")
    text = replace_once(text, "Adoption Status: NOT YET ADOPTED", "Constitutional Adoption Status: ADOPTED")
    text = replace_once(text, "Lock Status: NOT YET LOCKED", "Constitutional Lock Status: NOT YET LOCKED")
    text = replace_once(text, "Runtime Authority: FALSE\n\nProduction Authority: FALSE", "Runtime Authority: FALSE\n\nMigration Authority: FALSE\n\nProduction Authority: FALSE")
    write(PRIVACY_MD, text)

    document = Document(PRIVACY_DOCX)
    update_paragraph(document, "FOUNDER APPROVED - FOUNDER-APPROVED CANON SOURCE - NOT YET ADOPTED - NOT YET LOCKED", "FOUNDER APPROVED - CONSTITUTIONALLY ADOPTED - NOT YET LOCKED")
    update_paragraph(document, "Constitutional Status: FOUNDER-APPROVED CANON SOURCE", "Constitutional Status: FOUNDER-APPROVED ADOPTED CANON")
    update_paragraph(document, "Adoption Status: NOT YET ADOPTED", "Constitutional Adoption Status: ADOPTED")
    update_paragraph(document, "Lock Status: NOT YET LOCKED", "Constitutional Lock Status: NOT YET LOCKED")
    runtime = next(p for p in document.paragraphs if p.text.strip() == "Runtime Authority: FALSE")
    insert_after(runtime, "Migration Authority: FALSE")
    document.save(PRIVACY_DOCX)


def update_reporting() -> None:
    document = Document(REPORTING_DOCX)
    table = document.tables[0]
    values = [
        ("Founder Approval Status", "FOUNDER APPROVED"),
        ("Constitutional Adoption Status", "ADOPTED"),
        ("Constitutional Lock Status", "NOT YET LOCKED"),
        ("Implementation Authority", "FALSE"),
        ("Runtime Authority", "FALSE"),
        ("Migration Authority", "FALSE"),
        ("Provider Activation Authority", "FALSE"),
        ("Production Authority", "FALSE"),
        ("Public Launch Authority", "FALSE"),
        ("Public Trust Claim Authority", "FALSE"),
    ]
    while len(table.rows) < len(values):
        table.add_row()
    for row, pair in zip(table.rows, values):
        row.cells[0].text, row.cells[1].text = pair
    update_paragraph(document, "Constitutional Status: Founder Approved; Controlling Canon", "Constitutional Status: Founder Approved; Adopted; Controlling Canon; Not Yet Locked")
    implementation = next(p for p in document.paragraphs if p.text.strip() == "Implementation Authority: FALSE")
    insert_after(implementation, "Runtime Authority: FALSE")
    update_paragraph(document, "Schema or Migration Authority: FALSE", "Migration Authority: FALSE")
    update_paragraph(document, "External Vendor Activation: FALSE", "Provider Activation Authority: FALSE")
    update_paragraph(document, "Production-Data Access: FALSE", "Production Authority: FALSE")
    update_paragraph(document, "Public Reporting or Trust-Claim Authority: FALSE", "Public Trust Claim Authority: FALSE")
    document.save(REPORTING_DOCX)

    module_path = COMPLETION / "tools/build_completion.py"
    spec = importlib.util.spec_from_file_location("counterpart_builder", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader
    spec.loader.exec_module(module)
    extracted = module.docx_to_markdown(REPORTING_DOCX)
    wrapper = f"""# Reporting Format-Counterpart Control

| Field | Value |
|---|---|
| Founder Approval Status | FOUNDER APPROVED |
| Constitutional Status | FOUNDER-APPROVED ADOPTED CANON |
| Constitutional Adoption Status | ADOPTED |
| Constitutional Lock Status | NOT YET LOCKED |
| Implementation Authority | FALSE |
| Runtime Authority | FALSE |
| Migration Authority | FALSE |
| Production Authority | FALSE |
| Provider Activation Authority | FALSE |
| Public Launch Authority | FALSE |
| Public Trust Claim Authority | FALSE |
| Evidentiary Classification | FOUNDER_AUTHORIZED_DETERMINISTIC_FORMAT_COUNTERPART_DERIVED_FROM_EXACT_SUBSTANTIVE_SOURCE |

## Format-Counterpart Provenance

This Markdown file is a founder-authorized deterministic format counterpart generated from the exact mounted DOCX source with SHA-256 `d6cc3cc2add501d2d34d947de4e7a2ee7c6af77aa9669606c95514361657504d`. It is not represented as a recovered historical Markdown source. The preserved pre-adoption DOCX remains the exact substantive source evidence; the current DOCX differs only by founder-authorized document-control metadata.

# Exact Substantive Source Content

{extracted}
"""
    write(REPORTING_MD, wrapper)


def verify_substantive_immutability() -> dict[str, bool]:
    pre_privacy_md = (PRIVACY_DIR / "pre_adoption" / PRIVACY_MD.name).read_text(encoding="utf-8")
    post_privacy_md = PRIVACY_MD.read_text(encoding="utf-8")
    privacy_md_equal = pre_privacy_md.split("## Authority Boundary", 1)[1] == post_privacy_md.split("## Authority Boundary", 1)[1]

    pre_privacy_docx = Document(PRIVACY_DIR / "pre_adoption" / PRIVACY_DOCX.name)
    post_privacy_docx = Document(PRIVACY_DOCX)
    def body(document: Document, marker: str) -> list[str]:
        texts = [paragraph.text for paragraph in document.paragraphs]
        index = texts.index(marker)
        return texts[index:]
    privacy_docx_equal = body(pre_privacy_docx, "Authority Boundary") == body(post_privacy_docx, "Authority Boundary")

    pre_reporting = Document(REPORTING_DIR / "pre_adoption" / REPORTING_DOCX.name)
    post_reporting = Document(REPORTING_DOCX)
    reporting_docx_equal = body(pre_reporting, "1. Purpose") == body(post_reporting, "1. Purpose")
    return {
        "privacy_markdown_substantive_body_unchanged": privacy_md_equal,
        "privacy_docx_substantive_body_unchanged": privacy_docx_equal,
        "reporting_docx_substantive_body_unchanged": reporting_docx_equal,
    }


def adoption_record(row: str, title: str, source: Path, counterpart: Path, directory: Path, post_hashes: dict[Path, str]) -> Path:
    path = directory / ("C0_023_PRIVACY_AND_DATA_PROTECTION_CONSTITUTIONAL_ADOPTION_RECORD.md" if row == "C0-023" else "C0_035_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_CONSTITUTIONAL_ADOPTION_RECORD.md")
    pre_source = PRE_HASHES[source]
    pre_counterpart = PRE_HASHES[counterpart]
    text = f"""# {title} Constitutional Adoption Record

Disposition: `{'C0_023_PRIVACY_AND_DATA_PROTECTION_CONSTITUTIONALLY_ADOPTED' if row == 'C0-023' else 'C0_035_REPORTING_ANALYTICS_AND_BUSINESS_INTELLIGENCE_CONSTITUTIONALLY_ADOPTED'}`

- C0 identifier: `{row}`
- Canon title: `{title}`
- Controlling version: `2.0`
- Exact source-authority path: `{rel(source)}`
- Pre-adoption source-authority SHA-256: `{pre_source}`
- Post-adoption adopted artifact SHA-256: `{post_hashes[source]}`
- Format-counterpart path: `{rel(counterpart)}`
- Pre-adoption format-counterpart SHA-256: `{pre_counterpart}`
- Post-adoption format-counterpart SHA-256: `{post_hashes[counterpart]}`
- Founder-approval authority: `FOUNDER_AUTHORIZED_C0_023_AND_C0_035_CONSTITUTIONAL_ADOPTION`
- Founder event record: `{rel(ROOT / 'C0_023_C0_035_FOUNDER_ADOPTION_EVENT_RECORD.json')}`
- Source-reconciliation report: `{rel(COMPLETION / 'PRIVACY_REPORTING_FORMAT_COUNTERPART_COMPLETION_REPORT.md')}`
- Prerequisite evidence package: `{rel(PACKAGE)}`
- Prerequisite evidence-package SHA-256: `{PACKAGE_SHA}`
- Parity validation: `PASS`; omitted 0; added 0; changed requirement identifiers 0; changed authority boundaries 0; unexplained semantic deltas 0
- Visual QA: `PASS`; Privacy 24/24 pages; Reporting 15/15 pages
- Adoption-readiness findings: `P0 0; open P1 0; adoption-blocking P2 0`
- Retained P2 findings: `0`
- Adoption decision: `ADOPTED`
- Adoption timestamp: `{ADOPTION_TS}`
- Constitutional authority: `CONTROLLING WITHIN ASSIGNED DOMAIN`
- Supersession status: no predecessor or competing source is silently overwritten; pre-adoption artifacts remain preserved under `{rel(directory / 'pre_adoption')}`
- Amendment requirement: substantive amendment requires a separate governed founder directive, provenance preservation, constitutional review, and adoption decision
- Lock readiness: `NOT YET DETERMINED`
- Constitutional lock: `NOT YET LOCKED`
- Implementation authority: `FALSE`
- Runtime authority: `FALSE`
- Migration authority: `FALSE`
- Provider activation authority: `FALSE`
- Production authority: `FALSE`
- Public-launch authority: `FALSE`
- Public-trust-claim authority: `FALSE`
- Adoption-record SHA-256: recorded in the detached sidecar `{path.name}.sha256`

`ADOPTION DOES NOT CONSTITUTE CONSTITUTIONAL LOCK OR IMPLEMENTATION AUTHORITY`
"""
    write(path, text)
    digest = sha256(path)
    write(directory / f"{path.name}.sha256", f"{digest}  {path.name}")
    return path


def update_json_rows(path: Path, collection: str, updates: dict[str, dict]) -> None:
    data = json.loads(path.read_text(encoding="utf-8"))
    rows = data[collection]
    seen = set()
    for row in rows:
        row_id = row.get("record_id") or row.get("row_id")
        if row_id in updates:
            row.update(updates[row_id])
            seen.add(row_id)
    if seen != set(updates):
        raise ValueError(f"Missing rows in {path}: {set(updates) - seen}")
    if path.name == "C0_CURRENT_LIFECYCLE_STATE_REGISTER.json":
        data["category_counts"] = dict(Counter(row["current_category"] for row in rows))
    write_json(path, data)


def prepend_note(path: Path, note: str) -> None:
    text = path.read_text(encoding="utf-8")
    if note in text:
        return
    first_break = text.find("\n")
    write(path, text[: first_break + 1] + "\n" + note + "\n" + text[first_break + 1 :])


def update_governance(post_hashes: dict[Path, str], privacy_record: Path, reporting_record: Path) -> list[Path]:
    source_register = REPO / "docs/canon/reviews/governance_v1_0_c0_lifecycle_completion_batch_1/C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.json"
    lifecycle = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_CURRENT_LIFECYCLE_STATE_REGISTER.json"
    ledger = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.json"
    delta = REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_HISTORICAL_TO_CURRENT_DELTA.json"
    common = {
        "current_category": "lock_required",
        "current_exact_source_status": "EXACT_ADOPTED_SOURCE_VERIFIED",
        "founder_status": "FOUNDER_ADOPTED",
        "adoption_state": "ADOPTED",
        "lock_state": "NOT_YET_LOCKED",
        "lock_readiness_state": "NOT_YET_DETERMINED",
        "unresolved_blocker": True,
        "remedy": "Complete a separate checksum-backed founder lock-readiness review and explicit lock directive.",
        "retained_p2": [],
    }
    update_json_rows(source_register, "rows", {
        "C0-023": {**common, "current_repository_path": rel(PRIVACY_MD), "current_sha256": post_hashes[PRIVACY_MD], "current_controlling_artifact": rel(PRIVACY_MD), "lifecycle_evidence": rel(privacy_record)},
        "C0-035": {**common, "current_repository_path": rel(REPORTING_DOCX), "current_sha256": post_hashes[REPORTING_DOCX], "current_controlling_artifact": rel(REPORTING_DOCX), "lifecycle_evidence": rel(reporting_record)},
    })
    update_json_rows(lifecycle, "rows", {
        "C0-023": {**common, "current_repository_path": rel(PRIVACY_MD), "current_sha256": post_hashes[PRIVACY_MD], "current_controlling_artifact": rel(PRIVACY_MD), "lifecycle_evidence": rel(privacy_record)},
        "C0-035": {**common, "current_repository_path": rel(REPORTING_DOCX), "current_sha256": post_hashes[REPORTING_DOCX], "current_controlling_artifact": rel(REPORTING_DOCX), "lifecycle_evidence": rel(reporting_record)},
    })
    update_json_rows(ledger, "records", {
        "C0-023": {"current_category": "lock_required", "source": rel(PRIVACY_MD), "expected_c0_sha256": post_hashes[PRIVACY_MD], "founder_status": "FOUNDER_ADOPTED", "adoption_state": "ADOPTED", "lock_state": "NOT_YET_LOCKED", "next_step": "Separate founder lock-readiness review and explicit lock directive.", "adoption_record": rel(privacy_record)},
        "C0-035": {"current_category": "lock_required", "source": rel(REPORTING_DOCX), "expected_c0_sha256": post_hashes[REPORTING_DOCX], "founder_status": "FOUNDER_ADOPTED", "adoption_state": "ADOPTED", "lock_state": "NOT_YET_LOCKED", "next_step": "Separate founder lock-readiness review and explicit lock directive.", "adoption_record": rel(reporting_record)},
    })
    update_json_rows(delta, "rows", {
        "C0-023": {"current_category": "lock_required", "current_founder_status": "FOUNDER_ADOPTED", "current_adoption": "ADOPTED", "current_lock": "NOT_YET_LOCKED", "current_path": rel(PRIVACY_MD)},
        "C0-035": {"current_category": "lock_required", "current_founder_status": "FOUNDER_ADOPTED", "current_adoption": "ADOPTED", "current_lock": "NOT_YET_LOCKED", "current_path": rel(REPORTING_DOCX)},
    })

    md_paths = {
        "source": REPO / "docs/canon/reviews/governance_v1_0_c0_lifecycle_completion_batch_1/C0_CURRENT_SOURCE_OF_TRUTH_REGISTER.md",
        "lifecycle": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_CURRENT_LIFECYCLE_STATE_REGISTER.md",
        "ledger": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_ROW_BY_ROW_LIFECYCLE_RESOLUTION_LEDGER.md",
        "blocker": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_UNRESOLVED_LIFECYCLE_BLOCKER_LEDGER.md",
        "delta": REPO / "docs/canon/reviews/governance_v1_0_final_baseline_resumption/c0_lifecycle/C0_HISTORICAL_TO_CURRENT_DELTA.md",
    }
    rows = [
        ("C0-023", "Master Privacy and Data Protection Model", PRIVACY_MD, privacy_record),
        ("C0-035", "Master Reporting, Analytics, and Business Intelligence Model", REPORTING_DOCX, reporting_record),
    ]
    for row_id, family, artifact, record in rows:
        replacements = {
            "source": f"| {row_id} | {family} | lock_required | EXACT_ADOPTED_SOURCE_VERIFIED | FOUNDER_ADOPTED | ADOPTED | NOT_YET_LOCKED | `{rel(artifact)}` |",
            "lifecycle": f"| {row_id} | {family} | EXACT_ADOPTED_SOURCE_VERIFIED | ADOPTED | NOT_YET_LOCKED | `{rel(record)}` | True |",
            "ledger": f"| {row_id} | TRACK_D_MISSING_SOURCE_RECOVERY | lock_required | FOUNDER_ADOPTED | ADOPTED | NOT_YET_LOCKED | Separate founder lock-readiness review and explicit lock directive. |",
            "blocker": f"| {row_id} | {family} | lock_required | FOUNDER_ADOPTED | Adopted and controlling; constitutional lock remains separately gated. |",
            "delta": f"| {row_id} | UNRESOLVED | FOUNDER_ADOPTED | UNRESOLVED | ADOPTED | `{rel(artifact)}` |",
        }
        for name, replacement in replacements.items():
            text = md_paths[name].read_text(encoding="utf-8")
            text, count = re.subn(rf"^\| {row_id} \|.*$", replacement, text, flags=re.M)
            if count != 1:
                raise ValueError(f"Expected one {row_id} row in {md_paths[name]}, found {count}")
            write(md_paths[name], text)
    text = md_paths["ledger"].read_text(encoding="utf-8")
    text = re.sub(r"^Status:.*$", "Status: `UPDATED_AFTER_C0_023_C0_035_CONSTITUTIONAL_ADOPTION`; remaining lifecycle blockers: `17`. C0-023 and C0-035 are adopted and now remain blocked only on separately authorized constitutional lock.", text, count=1, flags=re.M)
    write(md_paths["ledger"], text)

    note = f"**July 16, 2026 C0-023/C0-035 adoption:** Privacy V2.0 and Reporting V2.0 are constitutionally `ADOPTED` and controlling within their assigned domains. Both remain `NOT YET LOCKED`; lock readiness is not yet determined. Adoption records: `{rel(privacy_record)}` and `{rel(reporting_record)}`. No implementation, runtime, migration, provider, production, launch, certification, or public-trust authority is created."
    companion_md = [
        REPO / "docs/canon/companions/EQUINESYNC_CONSTITUTIONAL_AUTHORITY_MATRIX_V1_2.md",
        REPO / "docs/canon/companions/CONSTITUTIONAL_DOMAIN_OWNERSHIP_AND_BOUNDARY_REGISTER_V1_1.md",
        REPO / "docs/canon/companions/EQUINESYNC_CONSTITUTIONAL_CROSS_REFERENCE_INDEX_V1_2.md",
        REPO / "docs/canon/companions/CROSS_CANON_REFERENCE_NORMALIZATION_REGISTER_V1_1.md",
        REPO / "docs/canon/companions/EQUINESYNC_CANON_DEPENDENCY_MAP_V1_2.md",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_FOUNDER_DECISION_REGISTER_V1_0.md",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_GOVERNANCE_REQUIREMENT_INDEX_V1_0.md",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_REQUIREMENT_TRACEABILITY_MATRIX_V1_0.md",
    ]
    for path in companion_md:
        prepend_note(path, note)

    event = {
        "disposition": "FOUNDER_AUTHORIZED_C0_023_AND_C0_035_CONSTITUTIONAL_ADOPTION",
        "timestamp": ADOPTION_TS,
        "rows": ["C0-023", "C0-035"],
        "adoption": True,
        "lock": False,
        "authority": AUTHORITY,
        "records": [rel(privacy_record), rel(reporting_record)],
    }
    for path in [
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_FOUNDER_DECISION_REGISTER_V1_0.json",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_GOVERNANCE_REQUIREMENT_INDEX_V1_0.json",
        REPO / "docs/canon/companions/MASTER_EQUINESYNC_REQUIREMENT_TRACEABILITY_MATRIX_V1_0.json",
    ]:
        data = json.loads(path.read_text(encoding="utf-8"))
        data["c0_023_c0_035_constitutional_adoption"] = event
        write_json(path, data)

    canon_index = REPO / "docs/canon/CANON_INDEX.md"
    text = canon_index.read_text(encoding="utf-8")
    text = re.sub(r"^\| 3 \| Master Privacy and Data Protection Model \|.*$", f"| 3 | Master Privacy and Data Protection Model | `{rel(PRIVACY_MD)}` | `ADOPTED`; controlling; not yet locked | Privacy and data-protection constitutional authority. No implementation, runtime, migration, provider, production, launch, certification, or public-trust authority. |", text, flags=re.M)
    text = re.sub(r"^\| 3 \| Master Reporting, Analytics, and Business Intelligence Model \|.*$", f"| 3 | Master Reporting, Analytics, and Business Intelligence Model | `{rel(REPORTING_DOCX)}` | `ADOPTED`; controlling; not yet locked | Reporting and analytics constitutional authority. No implementation, runtime, migration, provider, production, launch, certification, or public-trust authority. |", text, flags=re.M)
    write(canon_index, text)

    inventory = REPO / "docs/canon/registries/CANON_ARTIFACT_INVENTORY.md"
    text = inventory.read_text(encoding="utf-8")
    text = re.sub(r"^\| Master Privacy and Data Protection Model founder-approved Markdown \|.*$", f"| Master Privacy and Data Protection Model adopted Markdown | `{rel(PRIVACY_MD)}` | 2.0 | `{post_hashes[PRIVACY_MD]}` | `ADOPTED`; controlling; not yet locked | Pre-adoption hash `{PRE_HASHES[PRIVACY_MD]}` preserved; DOCX remains deterministic counterpart |", text, flags=re.M)
    text = re.sub(r"^\| Master Reporting, Analytics, and Business Intelligence Model founder-approved DOCX \|.*$", f"| Master Reporting, Analytics, and Business Intelligence Model adopted DOCX | `{rel(REPORTING_DOCX)}` | 2.0 | `{post_hashes[REPORTING_DOCX]}` | `ADOPTED`; controlling; not yet locked | Pre-adoption hash `{PRE_HASHES[REPORTING_DOCX]}` preserved; Markdown remains deterministic counterpart |", text, flags=re.M)
    write(inventory, text)

    state_registry = REPO / "docs/canon/registries/CANON_STATE_AND_LOCK_REGISTRY.md"
    text = state_registry.read_text(encoding="utf-8")
    additions = [
        f"| Privacy and Data Protection V2.0 | Adopted controlling constitutional canon | Founder adoption approved | Yes | `NOT YET LOCKED` | SHA-256 `{post_hashes[PRIVACY_MD]}`; P0/P1/blocking P2 0; all operational authority false |",
        f"| Reporting, Analytics, and Business Intelligence V2.0 | Adopted controlling constitutional canon | Founder adoption approved | Yes | `NOT YET LOCKED` | SHA-256 `{post_hashes[REPORTING_DOCX]}`; P0/P1/blocking P2 0; all operational authority false |",
    ]
    if additions[0] not in text:
        anchor = "| Record Governance Gap Matrix V2.0"
        index = text.find(anchor)
        text = text[:index] + "\n".join(additions) + "\n" + text[index:]
        write(state_registry, text)

    return [source_register, lifecycle, ledger, delta, *md_paths.values(), *companion_md, canon_index, inventory, state_registry]


def update_global_registers(post_hashes: dict[Path, str], privacy_record: Path, reporting_record: Path) -> list[Path]:
    paths = []
    adoption_md = REPO / "docs/canon/companions/GLOBAL_ADOPTION_STATE_REGISTER.md"
    text = adoption_md.read_text(encoding="utf-8")
    rows = [
        f"| Master Privacy and Data Protection Model V2.0 | adopted; not locked | {rel(privacy_record)} |",
        f"| Master Reporting, Analytics, and Business Intelligence Model V2.0 | adopted; not locked | {rel(reporting_record)} |",
    ]
    if rows[0] not in text:
        text += "\n" + "\n".join(rows) + "\n"
        write(adoption_md, text)
    paths.append(adoption_md)
    adoption_json = REPO / "docs/canon/companions/GLOBAL_ADOPTION_STATE_REGISTER.json"
    data = json.loads(adoption_json.read_text(encoding="utf-8"))
    data["entries"] = [x for x in data["entries"] if x["artifact"] not in {"Master Privacy and Data Protection Model V2.0", "Master Reporting, Analytics, and Business Intelligence Model V2.0"}]
    data["entries"].extend([
        {"artifact": "Master Privacy and Data Protection Model V2.0", "path": rel(PRIVACY_MD), "adoption_state": "adopted", "lock_state": "not yet locked", "evidence": rel(privacy_record), "evidence_sha256": sha256(privacy_record)},
        {"artifact": "Master Reporting, Analytics, and Business Intelligence Model V2.0", "path": rel(REPORTING_DOCX), "adoption_state": "adopted", "lock_state": "not yet locked", "evidence": rel(reporting_record), "evidence_sha256": sha256(reporting_record)},
    ])
    write_json(adoption_json, data); paths.append(adoption_json)

    lock_md = REPO / "docs/canon/companions/GLOBAL_LOCK_STATE_REGISTER.md"
    text = lock_md.read_text(encoding="utf-8")
    lock_rows = [
        f"| Master Privacy and Data Protection Model V2.0 | not yet locked | {rel(privacy_record)} |",
        f"| Master Reporting, Analytics, and Business Intelligence Model V2.0 | not yet locked | {rel(reporting_record)} |",
    ]
    if lock_rows[0] not in text:
        text += "\n" + "\n".join(lock_rows) + "\n"
        write(lock_md, text)
    paths.append(lock_md)
    lock_json = REPO / "docs/canon/companions/GLOBAL_LOCK_STATE_REGISTER.json"
    data = json.loads(lock_json.read_text(encoding="utf-8"))
    data["entries"] = [x for x in data["entries"] if x["artifact"] not in {"Master Privacy and Data Protection Model V2.0", "Master Reporting, Analytics, and Business Intelligence Model V2.0"}]
    data["entries"].extend([
        {"artifact": "Master Privacy and Data Protection Model V2.0", "path": rel(PRIVACY_MD), "lock_state": "not yet locked", "evidence": rel(privacy_record), "evidence_sha256": sha256(privacy_record)},
        {"artifact": "Master Reporting, Analytics, and Business Intelligence Model V2.0", "path": rel(REPORTING_DOCX), "lock_state": "not yet locked", "evidence": rel(reporting_record), "evidence_sha256": sha256(reporting_record)},
    ])
    write_json(lock_json, data); paths.append(lock_json)
    return paths


def main() -> None:
    preflight()
    preserve_pre_adoption()
    write_json(ROOT / "C0_023_C0_035_FOUNDER_ADOPTION_EVENT_RECORD.json", {
        "disposition": "FOUNDER_AUTHORIZED_C0_023_AND_C0_035_CONSTITUTIONAL_ADOPTION",
        "founder_decision": "APPROVED_AND_AUTHORIZED_FOR_ADOPTION",
        "timestamp": ADOPTION_TS,
        "prerequisite_disposition": "PRIVACY_AND_REPORTING_FOUNDER_APPROVED_FORMAT_COUNTERPARTS_COMPLETE_PENDING_SEPARATE_ADOPTION",
        "prerequisite_package": rel(PACKAGE),
        "prerequisite_package_sha256": PACKAGE_SHA,
        "rows": ["C0-023", "C0-035"],
        "authorization": "constitutional adoption only",
        "authority": AUTHORITY,
    })
    update_privacy()
    update_reporting()
    immutability = verify_substantive_immutability()
    if not all(immutability.values()):
        raise SystemExit(f"Substantive immutability failed: {immutability}")
    post_hashes = {path: sha256(path) for path in PRE_HASHES}
    privacy_record = adoption_record("C0-023", "Master Privacy and Data Protection Model V2.0", PRIVACY_MD, PRIVACY_DOCX, PRIVACY_DIR, post_hashes)
    reporting_record = adoption_record("C0-035", "Master Reporting, Analytics, and Business Intelligence Model V2.0", REPORTING_DOCX, REPORTING_MD, REPORTING_DIR, post_hashes)

    for directory, row in [(PRIVACY_DIR, "C0-023"), (REPORTING_DIR, "C0-035")]:
        write(directory / f"{row.replace('-', '_')}_RETAINED_P2_REGISTER.md", f"# {row} Retained P2 Register\n\nRetained P2 findings: `0`. Adoption-blocking P2 findings: `0`. Constitutional lock remains separately gated.\n")
        write(directory / f"{row.replace('-', '_')}_LOCK_READINESS_PREPARATION.md", f"# {row} Lock-Readiness Preparation\n\nStatus: `NOT YET DETERMINED`\n\nThe canon is adopted but not locked. A future lock-readiness review must independently reverify source identity, post-adoption hashes, cross-canon references, requirements, dependencies, retained P2 accounting, authority boundaries, and amendment/supersession posture. This preparation does not issue or authorize lock.\n")

    updated = update_governance(post_hashes, privacy_record, reporting_record)
    updated.extend(update_global_registers(post_hashes, privacy_record, reporting_record))
    write_json(ROOT / "C0_023_C0_035_ADOPTION_BUILD_RECORD.json", {
        "timestamp": ADOPTION_TS,
        "pre_adoption_hashes": {rel(path): digest for path, digest in PRE_HASHES.items()},
        "post_adoption_hashes": {rel(path): digest for path, digest in post_hashes.items()},
        "substantive_immutability": immutability,
        "adoption_records": [rel(privacy_record), rel(reporting_record)],
        "updated_governance_artifacts": sorted({rel(path) for path in updated}),
        "retained_p2": {"C0-023": 0, "C0-035": 0},
        "findings": {"p0": 0, "open_p1": 0, "adoption_blocking_p2": 0},
        "authority": AUTHORITY,
    })
    print(json.dumps({"timestamp": ADOPTION_TS, "post_hashes": {rel(p): h for p, h in post_hashes.items()}, "immutability": immutability}, indent=2))


if __name__ == "__main__":
    main()
